"""Tests for worksheet decomposition, solving, and formatting — migrated from run_tests.py."""

from __future__ import annotations

import asyncio
import json
import tempfile
from pathlib import Path

from schulpipeline.backends.base import LLMResponse
from schulpipeline.backends.router import BackendRouter
from schulpipeline.config import BackendConfig, OutputConfig, PipelineConfig, ResearchConfig
from schulpipeline.presets import resolve_preset, resolve_quick
from schulpipeline.worksheet import (
    DecomposeStage,
    SolveStage,
    format_worksheet_as_docx,
    format_worksheet_as_md,
)

# ---------------------------------------------------------------------------
# Mock infrastructure (duplicated from run_tests.py for standalone use)
# ---------------------------------------------------------------------------

DEFAULT_RESPONSES = {
    "Aufgaben-Parser für Schularbeitsblätter": json.dumps(
        {
            "title": "Gewinnverteilung OHG/KG",
            "subject": "Wirtschaft",
            "tasks": [
                {
                    "id": "aufgabe_1a",
                    "label": "Übung 2, Aufgabe 1a",
                    "task_type": "table_fill",
                    "text": "Der Gewinn der Baumeister OHG beträgt 130.000 €. Berechnen Sie die Gewinnaufteilung nach §121 HGB.",
                    "context": "4% Zinsen auf Kapitalanteil, Rest nach Köpfen",
                    "table_structure": {
                        "headers": ["Name", "Kapitalanteil", "4% Zinsen", "Restgewinn", "Gesamtgewinn"],
                        "rows": [
                            {"label": "Bauer", "values": ["100.000 €", "", "", ""]},
                            {"label": "Müller", "values": ["150.000 €", "", "", ""]},
                        ],
                    },
                    "data": {"gewinn": 130000, "zinssatz": 0.04},
                    "external_url": "",
                    "solvable": True,
                    "skip_reason": "",
                },
                {
                    "id": "uebung_1",
                    "label": "Übung 1",
                    "task_type": "external_link",
                    "text": "Quiz auf lernnetz24.de",
                    "context": "",
                    "table_structure": None,
                    "data": {},
                    "external_url": "https://www.lernnetz24.de/bwl/hinweise/99.html",
                    "solvable": False,
                    "skip_reason": "Externer Link — Quiz muss im Browser bearbeitet werden",
                },
            ],
        }
    ),
    "Aufgaben-Löser für Schulaufgaben": json.dumps(
        {
            "answer": "Bauer erhält 25.000 € Gesamtgewinn, Müller erhält 66.000 €",
            "table_data": {
                "headers": ["Name", "Kapitalanteil", "4% Zinsen", "Restgewinn", "Gesamtgewinn"],
                "rows": [
                    {"label": "Bauer", "values": ["100.000 €", "4.000 €", "56.000 €", "60.000 €"]},
                    {"label": "Müller", "values": ["150.000 €", "6.000 €", "56.000 €", "62.000 €"]},
                    {"label": "Summe", "values": ["250.000 €", "10.000 €", "112.000 €", "122.000 €"]},
                ],
            },
            "calculation_steps": [
                "4% von 100.000 = 4.000 €",
                "4% von 150.000 = 6.000 €",
                "Restgewinn: 130.000 - 10.000 = 120.000 €",
                "Pro Kopf: 120.000 / 2 = 60.000 €",
            ],
            "confidence": 0.95,
        }
    ),
}


SAMPLE_WORKSHEET_INPUT = """Übung 1:
😎Testen Sie sich: Quiz (Klick auf "Sofort üben)
https://www.lernnetz24.de/bwl/hinweise/99.html

Übung 2: Gewinnverteilung

Aufgabe 1 a) Der Gewinn der Baumeister OHG beträgt 130.000,00 €. Die Einlage von Herrn Bauer beläuft sich auf 100.000,00 €, die von Frau Müller auf 150.000,00 €. Die Verteilung soll nach altem Recht § 121 HGB erfolgen.
"Von dem Jahresgewinne gebührt jedem Gesellschafter zunächst ein Anteil in Höhe von vier vom Hundert seines Kapitalanteils."
Berechnen Sie die Gewinnaufteilung!
Name Kapitalanteil 4% Zinsen Restgewinn Gesamtgewinn
Bauer
Müller
Summe
"""


class MockBackend:
    def __init__(self):
        self.name = "mock"
        self.is_free = True
        self.supports_vision = True
        self.max_context = 128_000
        self.model = "mock-model"
        self.calls = []

    async def complete(self, messages, temperature=0.3, max_tokens=4096, response_format=None):
        self.calls.append(messages)
        system_msg = messages[0]["content"] if messages else ""
        for key, response in DEFAULT_RESPONSES.items():
            if key in system_msg:
                return LLMResponse(content=response, model="mock", backend_name="mock")
        return LLMResponse(content="{}", model="mock", backend_name="mock")

    async def complete_vision(self, messages, temperature=0.3, max_tokens=4096):
        return await self.complete(messages, temperature, max_tokens)

    async def close(self):
        pass


def _make_mock_router():
    """Create a mock router without needing real backends."""
    config = PipelineConfig(
        backends={"mock": BackendConfig(name="mock", api_key="test", enabled=True)},
        cascade={
            stage: ["mock"]
            for stage in [
                "intake",
                "plan",
                "research",
                "synthesize",
                "artifact",
                "agent_codegen",
                "decompose",
                "solve",
                "classify_docs",
                "fill_template",
                "audit",
                "classify_report",
                "amendments",
            ]
        },
        research=ResearchConfig(enabled=False, use_web=False),
        output=OutputConfig(dir=tempfile.mkdtemp(), default_format="pptx", language="de"),
    )

    router = BackendRouter.__new__(BackendRouter)
    router.config = config
    router._backends = {"mock": MockBackend()}
    router._cooldowns = {}
    router._call_counts = {"mock": 0}
    router._total_cost = 0.0
    router._stage_costs = {}
    router._stage_tokens = {}
    return config, router


# ---------------------------------------------------------------------------
# Decompose & Solve stage tests
# ---------------------------------------------------------------------------


def test_decompose_stage():
    """DecomposeStage parses messy input into structured tasks."""
    config, router = _make_mock_router()
    stage = DecomposeStage()
    assert stage.name == "decompose"

    context = {
        "intake": {
            "task_text": SAMPLE_WORKSHEET_INPUT,
            "subject": "Wirtschaft",
            "task_type": "worksheet",
        },
    }

    result = asyncio.run(stage.run(context, router, config))
    assert result.success, f"Decompose failed: {result.errors}"
    assert result.data["title"] == "Gewinnverteilung OHG/KG"
    assert len(result.data["tasks"]) == 2

    # First task should be solvable
    t0 = result.data["tasks"][0]
    assert t0["solvable"]
    assert t0["task_type"] == "table_fill"
    assert t0.get("table_structure") is not None

    # Second task should be unsolvable (external link)
    t1 = result.data["tasks"][1]
    assert not t1["solvable"]
    assert t1["task_type"] == "external_link"


def test_solve_stage():
    """SolveStage solves decomposed tasks and separates unsolvable ones."""
    config, router = _make_mock_router()
    stage = SolveStage()

    decomposed = json.loads(DEFAULT_RESPONSES["Aufgaben-Parser für Schularbeitsblätter"])
    context = {
        "intake": {"task_text": "test", "subject": "Wirtschaft", "task_type": "worksheet"},
        "decompose": decomposed,
    }

    result = asyncio.run(stage.run(context, router, config))
    assert result.success, f"Solve failed: {result.errors}"

    # Should have 1 solved task (aufgabe_1a) and 1 unsolvable (uebung_1)
    assert len(result.data["solved_tasks"]) == 1
    assert len(result.data["unsolvable_tasks"]) == 1

    solved = result.data["solved_tasks"][0]
    assert "table_data" in solved["solution"]
    assert "calculation_steps" in solved["solution"]
    assert len(solved["solution"]["calculation_steps"]) == 4


# ---------------------------------------------------------------------------
# Formatter tests
# ---------------------------------------------------------------------------


def test_format_worksheet_as_md():
    """Markdown formatter produces correct structure."""
    solve_result = {
        "title": "Gewinnverteilung",
        "subject": "Wirtschaft",
        "solved_tasks": [
            {
                "task": {"label": "Aufgabe 1a", "text": "Berechnen Sie..."},
                "solution": {
                    "answer": "Bauer: 60.000 €, Müller: 62.000 €",
                    "table_data": {
                        "headers": ["Name", "Kapital", "Zinsen", "Rest", "Gesamt"],
                        "rows": [
                            {"label": "Bauer", "values": ["100.000 €", "4.000 €", "56.000 €", "60.000 €"]},
                            {"label": "Müller", "values": ["150.000 €", "6.000 €", "56.000 €", "62.000 €"]},
                        ],
                    },
                    "calculation_steps": ["4% von 100.000 = 4.000 €"],
                    "confidence": 0.95,
                },
            }
        ],
        "unsolvable_tasks": [
            {"label": "Übung 1", "skip_reason": "Externer Link", "external_url": "https://example.com"}
        ],
    }

    md = format_worksheet_as_md(solve_result)
    assert "# Gewinnverteilung" in md
    assert "Wirtschaft" in md
    assert "| Bauer |" in md
    assert "4.000 €" in md
    assert "Rechenweg" in md
    assert "Nicht bearbeitbare Aufgaben" in md
    assert "Externer Link" in md


def test_format_worksheet_as_docx():
    """DOCX formatter produces a valid file with tables."""
    solve_result = {
        "title": "Testblatt",
        "subject": "BWL",
        "solved_tasks": [
            {
                "task": {"label": "Aufgabe 1", "text": "Berechnen Sie den Gewinn."},
                "solution": {
                    "answer": "Ergebnis: 100.000 €",
                    "table_data": {
                        "headers": ["Name", "Betrag"],
                        "rows": [
                            {"label": "A", "values": ["50.000 €"]},
                            {"label": "B", "values": ["50.000 €"]},
                        ],
                    },
                    "calculation_steps": ["50.000 + 50.000 = 100.000"],
                    "confidence": 1.0,
                },
            }
        ],
        "unsolvable_tasks": [],
    }

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name

    try:
        format_worksheet_as_docx(solve_result, path)
        assert Path(path).exists()
        assert Path(path).stat().st_size > 0

        # Verify it's a valid docx
        from docx import Document

        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Testblatt" in text
        assert "Aufgabe 1" in text
        assert "100.000 €" in text

        # Verify table exists
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "Name"
        assert table.rows[1].cells[0].text == "A"
    finally:
        Path(path).unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# Preset & pipeline integration tests
# ---------------------------------------------------------------------------


def test_worksheet_preset_resolves():
    """Worksheet presets resolve correctly."""
    preset = resolve_preset("arbeitsblatt", "wirtschaft")
    assert preset.output_format == "worksheet"
    assert preset.output_constraints.get("worksheet_mode")
    assert preset.output_constraints.get("show_calculation_steps")
    assert preset.output_constraints.get("german_number_format")


def test_worksheet_quick_presets():
    """All FIAE-Blatt quick presets resolve."""
    for key in ["fiae-blatt-wirtschaft", "fiae-blatt-itsec", "fiae-blatt-prog"]:
        preset = resolve_quick(key)
        assert preset is not None, f"Quick preset {key} should resolve"
        assert preset.output_format == "worksheet"


def test_pipeline_selects_worksheet_stages():
    """Pipeline uses decompose -> solve flow for worksheet presets."""
    from schulpipeline.pipeline import Pipeline

    config, router = _make_mock_router()
    pipeline = Pipeline(config, router)

    preset = resolve_preset("arbeitsblatt", "wirtschaft")
    context = {"preset": preset}

    stages = pipeline._select_stages(context)
    stage_names = [s.name for s in stages]
    assert stage_names == ["intake", "decompose", "solve"]


def test_pipeline_standard_stages_for_presentation():
    """Pipeline still uses standard stages for non-worksheet presets."""
    from schulpipeline.pipeline import Pipeline

    config, router = _make_mock_router()
    pipeline = Pipeline(config, router)

    preset = resolve_preset("praesentation", "it_sicherheit")
    context = {"preset": preset}

    stages = pipeline._select_stages(context)
    stage_names = [s.name for s in stages]
    assert stage_names == ["intake", "plan", "research", "synthesize", "artifact"]
