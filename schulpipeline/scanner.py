"""Scan school document directories and classify files.

Handles the reality of school file dumps:
- Tasks mixed with info material in the same folder
- OneNote exports that embed other files + student answers
- PDF/PNG duplicates of .docx files
- Empty placeholder files
- Multi-document tasks spread across files
- Source code projects with .java, .drawio, .db, .zip

Classification is heuristic-first (free, fast), with optional LLM refinement.
"""

from __future__ import annotations

import json
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable

logger = logging.getLogger(__name__)


# ============================================================
# Data Model
# ============================================================


@dataclass
class ContentResult:
    """Standardised output from any file reader."""

    full_text: str = ""
    text_preview: str = ""  # first 500 chars
    paragraph_count: int = 0
    table_count: int = 0
    empty_table_cells: int = 0
    image_count: int = 0
    language: str = ""  # "java", "sql", "xml", "german", ""
    structure: dict = field(default_factory=dict)  # type-specific metadata
    error: str | None = None


@dataclass
class ScannedFile:
    """A single file with classification metadata."""

    path: str
    filename: str
    subject_folder: str
    size_bytes: int
    content_type: str

    # Classification
    role: str = "unknown"
    confidence: float = 0.0
    classification_method: str = "none"

    # Content analysis
    has_tables: bool = False
    has_images: bool = False
    paragraph_count: int = 0
    table_count: int = 0
    empty_table_cells: int = 0
    embedded_refs: list[str] = field(default_factory=list)
    task_signals: list[str] = field(default_factory=list)
    info_signals: list[str] = field(default_factory=list)
    text_preview: str = ""
    full_text: str = ""
    task_score: int = 0
    info_score: int = 0

    # Relationships
    duplicate_of: str | None = None
    related_files: list[str] = field(default_factory=list)
    bundle_id: str | None = None
    notes: str = ""


@dataclass
class TaskBundle:
    """A group of related files that form one logical task."""

    bundle_id: str
    title: str
    subject: str
    task_files: list[str] = field(default_factory=list)
    info_files: list[str] = field(default_factory=list)
    answer_files: list[str] = field(default_factory=list)
    duplicates: list[str] = field(default_factory=list)
    task_type: str = "unknown"
    notes: str = ""


@dataclass
class ScanResult:
    """Complete scan result for a directory."""

    scan_root: str
    total_files: int = 0
    files: list[ScannedFile] = field(default_factory=list)
    bundles: list[TaskBundle] = field(default_factory=list)
    subject_folders: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


# ============================================================
# Heuristic Signals
# ============================================================

TASK_KEYWORDS_STRONG = [
    r"\bAufgabe[n]?\b",
    r"\bÜbung(?:en)?\b",
    r"\bArbeitsteil\b",
    r"\bBearbeiten Sie\b",
    r"\bFormulieren Sie\b",
    r"\bBeantworten Sie\b",
    r"\bRecherchieren Sie\b",
    r"\bErstellen Sie\b",
    r"\bBeschreiben Sie\b",
    r"\bErklären Sie\b",
    r"\bBegründen Sie\b",
    r"\bNennen Sie\b",
    r"\bDefinieren Sie\b",
    r"\bVergleichen Sie\b",
    r"\bAnalysieren Sie\b",
    r"\bBewerten Sie\b",
    r"\bErörtern Sie\b",
    r"\bKreuzen Sie\b",
    r"\bErgänzen Sie\b",
    r"\bLesen Sie\b",
    r"\bSchreiben Sie\b",
    r"\bStellen Sie\b.*\bvor\b",
    r"\bSuchen Sie\b",
    r"\bÜberprüfen Sie\b",
    r"\bSammeln Sie\b",
    r"\bNotieren Sie\b",
]
TASK_KEYWORDS_WEAK = [
    r"\bLösung(?:en|sword)?\b",
    r"\bSilbenrätsel\b",
    r"\bSituation:\b",
    r"\bBearbeitungsvorschlag\b",
    r"\bKreuzen Sie an\b",
]
INFO_KEYWORDS = [
    r"\bInformation(?:steil)?\b",
    r"\bInfo:\b",
    r"\bDefinition\b",
    r"\bGrundsätzlich\b",
    r"\bMan unterscheidet\b",
    r"\bZu den\b.*\bzählen\b",
    r"\bBeispiel(?:e)?\b",
]

ONENOTE_EMBED_RE = re.compile(r"<<(.+?\.(docx|pdf|pptx))>>")
ONENOTE_FILENAME_RE = re.compile(r".*OneNote\.docx$", re.IGNORECASE)
DOCX_PDF_DUP_RE = re.compile(r"^(.+\.docx)\.pdf$")
TASK_FILENAME_RE = re.compile(r"(?:Aufgabe|Übung|Arbeit|_A\d+_)", re.IGNORECASE)
INFO_FILENAME_RE = re.compile(
    r"(?:_Info_|_Info\b|Information|_Texte?\b|_Text\b)",
    re.IGNORECASE,
)
TEXT_MATERIAL_FILENAME_RE = re.compile(
    r"(?:_Texte?\.docx$|_Text\.docx$)",
    re.IGNORECASE,
)
PLANNING_FILENAME_RE = re.compile(r"(?:Planung|Bewertung|Gliederung)", re.IGNORECASE)

# Roles
ROLES = {
    "task",
    "info",
    "answer",
    "onenote_export",
    "duplicate",
    "empty",
    "resource",
    "unknown",
}

# ============================================================
# Content Type Mapping
# ============================================================

TYPE_MAP: dict[str, str] = {
    # Documents
    ".docx": "docx",
    ".pdf": "pdf",
    ".txt": "txt",
    ".pptx": "pptx",
    ".xlsx": "xlsx",
    # Source code
    ".java": "java",
    ".py": "python",
    ".cs": "csharp",
    ".js": "javascript",
    ".html": "html",
    ".css": "css",
    ".fxml": "fxml",
    # Data
    ".db": "db",
    ".sqlite": "db",
    ".sql": "sql",
    ".csv": "csv",
    ".json": "json",
    # Diagrams
    ".drawio": "drawio",
    # Archives
    ".zip": "zip",
    # Images
    ".png": "png",
    ".jpg": "png",
    ".jpeg": "png",
    ".gif": "png",
}


# ============================================================
# Readers
# ============================================================


def read_docx(path: Path) -> ContentResult:
    """Read a .docx file. Extract paragraphs, tables, images."""
    try:
        from docx import Document

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables, total_empty = [], 0
        for t in doc.tables:
            rows = []
            for row in t.rows:
                cells = [c.text.strip()[:200] for c in row.cells]
                rows.append(cells)
                total_empty += sum(1 for c in cells if not c)
            tables.append(rows)
        image_count = len(doc.inline_shapes) if hasattr(doc, "inline_shapes") else 0
        full_text = "\n".join(paragraphs)
        return ContentResult(
            full_text=full_text,
            text_preview=full_text[:500],
            paragraph_count=len(paragraphs),
            table_count=len(tables),
            empty_table_cells=total_empty,
            image_count=image_count,
            structure={"paragraphs": paragraphs, "tables": tables},
        )
    except Exception as e:
        return ContentResult(error=str(e))


def read_txt(path: Path) -> ContentResult:
    """Read a plain text file."""
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        lines = [line for line in text.splitlines() if line.strip()]
        return ContentResult(
            full_text=text,
            text_preview=text[:500],
            paragraph_count=len(lines),
        )
    except Exception as e:
        return ContentResult(error=str(e))


def read_java(path: Path) -> ContentResult:
    """Read a Java (or other source) file. Extract structural metadata."""
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return ContentResult(error=str(e))

    classes = re.findall(
        r"(?:public\s+|abstract\s+|final\s+)*(?:class|interface|enum)\s+(\w+)",
        text,
    )
    methods = re.findall(
        r"(?:public|private|protected)\s+(?:static\s+)?(?:\w+\s+)+(\w+)\s*\(",
        text,
    )
    imports = re.findall(r"^import\s+(.+);", text, re.MULTILINE)
    has_main = bool(re.search(r"public\s+static\s+void\s+main\s*\(", text))
    loc = sum(1 for line in text.splitlines() if line.strip())
    is_template = bool(
        re.search(
            r"(?:TODO|FIXME|_NN|YOUR[\s_]CODE[\s_]HERE|//\s*\.\.\.|/\*\s*\.\.\.\s*\*/)",
            text,
            re.IGNORECASE,
        )
    )

    structure = {
        "classes": classes,
        "methods": methods,
        "imports": imports,
        "has_main": has_main,
        "loc": loc,
        "is_template": is_template,
    }

    # Detect language from extension
    suffix = path.suffix.lower()
    lang_map = {
        ".java": "java",
        ".py": "python",
        ".cs": "csharp",
        ".js": "javascript",
        ".html": "html",
        ".css": "css",
        ".fxml": "fxml",
    }
    language = lang_map.get(suffix, "code")

    return ContentResult(
        full_text=text,
        text_preview=text[:500],
        paragraph_count=loc,
        language=language,
        structure=structure,
    )


def read_drawio(path: Path) -> ContentResult:
    """Read a .drawio file (XML). Extract text labels and diagram type."""
    try:
        import xml.etree.ElementTree as ET

        tree = ET.parse(str(path))
        root = tree.getroot()

        labels = []
        for cell in root.iter("mxCell"):
            value = cell.get("value", "").strip()
            if value and not value.startswith("<"):
                labels.append(value)

        all_cells = list(root.iter("mxCell"))
        structure = {
            "labels": labels,
            "cell_count": len(all_cells),
            "has_edges": any(c.get("edge") for c in all_cells),
            "diagram_name": root.get("name", path.stem),
        }

        full_text = "\n".join(labels)
        return ContentResult(
            full_text=full_text,
            text_preview=full_text[:500],
            paragraph_count=len(labels),
            language="xml",
            structure=structure,
        )
    except Exception as e:
        return ContentResult(error=str(e))


def read_sqlite(path: Path) -> ContentResult:
    """Inspect a SQLite database. Extract schema and row counts."""
    try:
        import sqlite3

        conn = sqlite3.connect(f"file:{path}?mode=ro", uri=True)
        cursor = conn.cursor()

        tables = []
        for (name,) in cursor.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'",
        ).fetchall():
            cols = cursor.execute(f"PRAGMA table_info('{name}')").fetchall()
            count = cursor.execute(f"SELECT COUNT(*) FROM '{name}'").fetchone()[0]
            tables.append(
                {
                    "name": name,
                    "columns": [{"name": c[1], "type": c[2]} for c in cols],
                    "row_count": count,
                }
            )

        conn.close()

        text_parts = [f"Database: {path.name}"]
        for t in tables:
            cols_str = ", ".join(f"{c['name']} {c['type']}" for c in t["columns"])
            text_parts.append(f"Table {t['name']} ({t['row_count']} rows): {cols_str}")

        full_text = "\n".join(text_parts)
        return ContentResult(
            full_text=full_text,
            text_preview=full_text[:500],
            table_count=len(tables),
            paragraph_count=len(tables),
            language="sql",
            structure={"tables": tables, "total_rows": sum(t["row_count"] for t in tables)},
        )
    except Exception as e:
        return ContentResult(error=str(e))


def read_zip(path: Path) -> ContentResult:
    """List ZIP archive contents and detect project type."""
    try:
        import zipfile

        if not zipfile.is_zipfile(str(path)):
            return ContentResult(error=f"Not a valid ZIP: {path.name}")

        with zipfile.ZipFile(str(path), "r") as zf:
            entries = zf.namelist()

        extensions = [Path(e).suffix.lower() for e in entries if Path(e).suffix]

        project_type = "unknown"
        if any(e.endswith(".java") for e in entries):
            project_type = "java"
            if any(e.endswith(".fxml") or e.endswith(".css") for e in entries):
                project_type = "javafx"
        elif any(e.endswith(".py") for e in entries):
            project_type = "python"

        structure = {
            "entries": entries[:50],
            "total_files": len(entries),
            "extensions": sorted(set(extensions)),
            "project_type": project_type,
            "has_src": any("src/" in e or "src\\" in e for e in entries),
        }

        full_text = f"Archive: {path.name}\n" + "\n".join(entries[:50])
        return ContentResult(
            full_text=full_text,
            text_preview=full_text[:500],
            paragraph_count=len(entries),
            structure=structure,
        )
    except Exception as e:
        return ContentResult(error=str(e))


def read_pdf(path: Path) -> ContentResult:
    """Extract text from a PDF file."""
    try:
        import pdfplumber

        with pdfplumber.open(str(path)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
            tables = []
            for p in pdf.pages:
                tables.extend(p.extract_tables() or [])
        full_text = "\n".join(pages)
        return ContentResult(
            full_text=full_text,
            text_preview=full_text[:500],
            paragraph_count=full_text.count("\n"),
            table_count=len(tables),
            structure={"page_count": len(pages)},
        )
    except ImportError:
        pass

    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        pages = [page.get_text() for page in doc]
        full_text = "\n".join(pages)
        doc.close()
        return ContentResult(
            full_text=full_text,
            text_preview=full_text[:500],
            paragraph_count=full_text.count("\n"),
            structure={"page_count": len(pages)},
        )
    except ImportError:
        pass

    # No PDF library available
    try:
        size = path.stat().st_size
    except OSError:
        size = 0
    return ContentResult(
        text_preview=(f"[PDF: {path.name}, {size} bytes — install pdfplumber for text extraction]"),
        structure={"page_count": 0},
        error="no_pdf_library",
    )


# Reader registry: maps content_type → reader function
READERS: dict[str, Callable[[Path], ContentResult]] = {
    "docx": read_docx,
    "txt": read_txt,
    "java": read_java,
    "python": read_java,  # same structural reader
    "csharp": read_java,
    "javascript": read_java,
    "html": read_java,
    "css": read_java,
    "fxml": read_java,
    "drawio": read_drawio,
    "db": read_sqlite,
    "zip": read_zip,
    "pdf": read_pdf,
}


# ============================================================
# Classification
# ============================================================


def _match_keywords(text: str, patterns: list[str]) -> list[str]:
    """:param text: The input text to search within.
    :type text: str
    :param patterns: A list of regex patterns to search for in the text.
    :type patterns: list[str]
    :return: A list of patterns that match the text.
    :rtype: list[str]
    """
    return [p for p in patterns if re.search(p, text, re.IGNORECASE)]


def classify_file(path: str | Path, scan_root: str | Path) -> ScannedFile:
    """Classify a single file using heuristics."""
    path, scan_root = Path(path), Path(scan_root)
    rel_path = str(path.relative_to(scan_root))
    filename = path.name
    # Walk up to find the subject folder (first child of scan_root)
    subject_folder = path.parent.name if path.parent != scan_root else ""
    if path.parent.parent != scan_root and path.parent.parent.is_dir():
        # Two levels deep — subject is the grandparent
        subject_folder = path.parent.parent.name
    size = path.stat().st_size
    suffix = path.suffix.lower()
    content_type = TYPE_MAP.get(suffix, "unknown")

    sf = ScannedFile(
        path=rel_path,
        filename=filename,
        subject_folder=subject_folder,
        size_bytes=size,
        content_type=content_type,
    )

    # --- Quick exits ---

    if size == 0:
        sf.role, sf.confidence, sf.classification_method = "empty", 1.0, "heuristic"
        return sf

    dup_match = DOCX_PDF_DUP_RE.match(filename)
    if dup_match:
        sf.role = "duplicate"
        sf.duplicate_of = str(Path(rel_path).parent / dup_match.group(1))
        sf.confidence, sf.classification_method = 0.95, "heuristic"
        return sf

    # --- Resource types (always resource role) ---

    if content_type == "db":
        reader = READERS["db"]
        cr = reader(path)
        _apply_content_result(sf, cr)
        sf.role, sf.confidence, sf.classification_method = "resource", 0.90, "heuristic"
        return sf

    if content_type == "zip":
        reader = READERS["zip"]
        cr = reader(path)
        _apply_content_result(sf, cr)
        # _NN.zip → task template, else resource
        if re.search(r"_NN\b", filename, re.IGNORECASE):
            sf.role, sf.confidence = "task", 0.85
        else:
            sf.role, sf.confidence = "resource", 0.80
        sf.classification_method = "heuristic"
        return sf

    if content_type == "drawio":
        reader = READERS["drawio"]
        cr = reader(path)
        _apply_content_result(sf, cr)
        sf.role, sf.confidence, sf.classification_method = "info", 0.80, "heuristic"
        return sf

    # --- Source code ---

    if content_type in ("java", "python", "csharp", "javascript", "html", "css", "fxml"):
        reader = READERS[content_type]
        cr = reader(path)
        _apply_content_result(sf, cr)
        structure = cr.structure
        # Classify based on filename and content signals
        if re.search(r"_NN\b", filename, re.IGNORECASE) or structure.get("is_template"):
            sf.role, sf.confidence = "task", 0.85
        elif re.search(r"(?:Lösung|Loesung|Solution)", filename, re.IGNORECASE):
            sf.role, sf.confidence = "info", 0.85
        elif structure.get("has_main"):
            sf.role, sf.confidence = "info", 0.70
        else:
            sf.role, sf.confidence = "resource", 0.60
        sf.classification_method = "heuristic"
        return sf

    # --- Images (no text to extract) ---

    if content_type == "png":
        # Check if in IMG/ subfolder → resource
        if "IMG" in Path(rel_path).parts or "img" in Path(rel_path).parts:
            sf.role, sf.confidence = "resource", 0.80
        else:
            sf.role, sf.confidence = "info", 0.50
        sf.classification_method = "heuristic"
        return sf

    # --- PDF (try text extraction, fall back to filename heuristics) ---

    if content_type == "pdf":
        reader = READERS["pdf"]
        cr = reader(path)
        _apply_content_result(sf, cr)
        if cr.error == "no_pdf_library" or not cr.full_text.strip():
            # No text extracted — use filename heuristics
            if TASK_FILENAME_RE.search(filename):
                sf.role, sf.confidence = "task", 0.60
            else:
                sf.role, sf.confidence = "info", 0.50
            sf.classification_method = "heuristic"
            return sf
        # Got text — fall through to keyword classification below

    # --- Text-based classification (docx, txt, pdf with text) ---

    if content_type in ("docx", "txt") or (content_type == "pdf" and sf.full_text):
        if content_type in ("docx", "txt"):
            reader = READERS.get(content_type)
            if reader:
                cr = reader(path)
                _apply_content_result(sf, cr)

        if not sf.full_text and not sf.text_preview:
            sf.role, sf.confidence, sf.classification_method = "unknown", 0.0, "heuristic"
            return sf

        full_text = sf.full_text

        # OneNote detection
        embeds = [m.group(1) for m in ONENOTE_EMBED_RE.finditer(full_text)]
        sf.embedded_refs = embeds

        if ONENOTE_FILENAME_RE.match(filename):
            cleaned = ONENOTE_EMBED_RE.sub("", full_text).strip()
            substantive = [
                line for line in cleaned.splitlines() if line.strip() and line.strip().lower() not in ("x", "fehlt", "")
            ]
            if len(substantive) > 3:
                sf.role, sf.confidence = "answer", 0.85
            else:
                sf.role, sf.confidence = "onenote_export", 0.95
            sf.classification_method = "heuristic"
            return sf

        # Keyword scoring
        strong = _match_keywords(full_text, TASK_KEYWORDS_STRONG)
        weak = _match_keywords(full_text, TASK_KEYWORDS_WEAK)
        info = _match_keywords(full_text, INFO_KEYWORDS)
        sf.task_signals, sf.info_signals = strong + weak, info

        task_score = len(strong) * 2 + len(weak)
        info_score = len(info)

        if TASK_FILENAME_RE.search(filename):
            task_score += 3
        if INFO_FILENAME_RE.search(filename):
            info_score += 3
        if TEXT_MATERIAL_FILENAME_RE.search(filename):
            info_score += 5
            task_score = max(0, task_score - 3)
        if PLANNING_FILENAME_RE.search(filename):
            info_score += 2
        if sf.empty_table_cells > 2:
            task_score += 3

        sf.task_score, sf.info_score = task_score, info_score

        if task_score >= 4 and task_score > info_score:
            sf.role = "task"
            sf.confidence = min(0.95, 0.5 + task_score * 0.05)
        elif info_score >= 3 and info_score >= task_score:
            sf.role = "info"
            sf.confidence = min(0.90, 0.5 + info_score * 0.05)
        elif task_score > 0 and task_score > info_score:
            sf.role = "task"
            sf.confidence = 0.4 + task_score * 0.05
        elif info_score > 0:
            sf.role = "info"
            sf.confidence = 0.4 + info_score * 0.05
        else:
            if sf.has_tables and sf.empty_table_cells > 2 and sf.paragraph_count < 15:
                sf.role, sf.confidence = "task", 0.45
            elif sf.paragraph_count > 20:
                sf.role, sf.confidence = "info", 0.4
            else:
                sf.role, sf.confidence = "unknown", 0.2

        if sf.paragraph_count > 80 and task_score >= 3 and info_score >= 3:
            sf.notes = "hybrid: info + embedded tasks"

        sf.classification_method = "heuristic"
        return sf

    # Fallback for unknown types
    sf.role, sf.confidence, sf.classification_method = "unknown", 0.1, "heuristic"
    return sf


def _apply_content_result(sf: ScannedFile, cr: ContentResult) -> None:
    """Copy ContentResult fields into a ScannedFile."""
    sf.text_preview = cr.text_preview
    sf.full_text = cr.full_text
    sf.paragraph_count = cr.paragraph_count
    sf.table_count = cr.table_count
    sf.has_tables = cr.table_count > 0
    sf.has_images = cr.image_count > 0
    sf.empty_table_cells = cr.empty_table_cells


# ============================================================
# Bundle Detection
# ============================================================


def _normalize_stem(fn: str) -> str:
    """Normalizes the stem of a file path by converting it to lowercase and removing specific suffixes and non-alphanumeric characters.

    :param fn: The file path whose stem needs normalization.
    :type fn: str
    :return: The normalized stem of the file path.
    :rtype: str
    """
    s = Path(fn).stem.lower()
    s = re.sub(r"onenote$", "", s)
    return re.sub(r"[^a-z0-9äöüß]", "", s)


def _stems_related(a: str, b: str) -> bool:
    """Determines if two strings are related based on a similarity threshold.

    :param a: First string to compare.
    :type a: str
    :param b: Second string to compare.
    :type b: str
    :return: True if the strings are related, False otherwise.
    :rtype: bool
    """
    if not a or not b:
        return False
    if a in b or b in a:
        return True
    common = sum(1 for ca, cb in zip(a, b) if ca == cb)
    return common >= min(len(a), len(b)) * 0.6


def _make_id(folder: str, fn: str) -> str:
    """Generates a unique identifier for a file based on its folder and name.

    :param folder: The folder containing the file.
    :type folder: str
    :param fn: The filename of the file.
    :type fn: str
    :return: A lowercase, URL-friendly string representing the file's identifier.
    :rtype: str
    """
    stem = re.sub(r"OneNote$", "", Path(fn).stem)
    safe = re.sub(r"_+", "_", re.sub(r"[^a-zA-Z0-9_äöüÄÖÜß-]", "_", stem)).strip("_")
    return f"{folder}_{safe}".lower()[:60]


def _title(sf: ScannedFile) -> str:
    """Returns the title of the file based on its text preview or filename.

    :param sf: The scanned file object.
    :type sf: ScannedFile
    :return: The title of the file.
    :rtype: str
    """
    if sf.text_preview:
        first = sf.text_preview.split("\n")[0].strip()
        if 5 < len(first) < 100 and not first.startswith("HOT "):
            return first
    return Path(sf.filename).stem.replace("_", " ").replace("OneNote", "").strip()


def _find_subject_folder(dir_path: str, scan_root: Path) -> str:
    """Walk up from dir_path to find the subject folder (first child of scan_root)."""
    p = Path(dir_path)
    while p.parent != scan_root and p.parent != p:
        p = p.parent
    return p.name if p != scan_root else ""


def detect_project_bundles(
    files: list[ScannedFile],
    scan_root: Path,
) -> list[TaskBundle]:
    """Detect directory-based project bundles.

    Heuristics:
    1. Directory contains mixed file types (.java + .docx + .drawio) → project
    2. Directory has LVL/Level progression in filenames → progressive task
    3. Directory contains _NN or Lösung files → task + solution pair
    4. IMG/ subdirectory → visual resources for parent
    """
    by_dir: dict[str, list[ScannedFile]] = {}
    for f in files:
        parent = str(Path(f.path).parent)
        by_dir.setdefault(parent, []).append(f)

    project_bundles = []
    for dir_path, dir_files in by_dir.items():
        types = {f.content_type for f in dir_files}

        # Mixed types (code + docs + diagrams) = project
        code_types = types & {"java", "python", "csharp", "javascript", "fxml"}
        doc_types = types & {"docx", "pdf", "txt"}
        other_types = types & {"drawio", "zip", "db"}
        is_project = bool(code_types) and bool(doc_types | other_types)

        # LVL progression detection
        has_levels = any(re.search(r"LVL\d|Level\d", f.filename, re.I) for f in dir_files)

        if not (is_project or has_levels):
            continue

        bundle = TaskBundle(
            bundle_id=_make_id(
                _find_subject_folder(dir_path, scan_root),
                Path(dir_path).name + "_project",
            ),
            title=Path(dir_path).name,
            subject=_find_subject_folder(dir_path, scan_root),
            task_type="project",
        )
        for f in dir_files:
            if f.role == "task" or re.search(r"_NN\b", f.filename, re.I):
                bundle.task_files.append(f.path)
            elif re.search(r"(?:Lösung|Loesung|Solution)", f.filename, re.I):
                bundle.answer_files.append(f.path)
            elif f.role == "resource" or f.content_type in ("db", "zip", "png"):
                bundle.duplicates.append(f.path)  # reuse as "resources" slot
            else:
                bundle.info_files.append(f.path)
        project_bundles.append(bundle)

    return project_bundles


def build_bundles(files: list[ScannedFile], scan_root: Path | None = None) -> list[TaskBundle]:
    """Build task bundles from classified files."""
    bundles: list[TaskBundle] = []
    used: set[str] = set()

    # First: detect project-level bundles if scan_root is available
    if scan_root:
        project_bundles = detect_project_bundles(files, scan_root)
        for pb in project_bundles:
            all_paths = pb.task_files + pb.info_files + pb.answer_files + pb.duplicates
            used.update(all_paths)
            bundles.append(pb)

    # Then: standard per-folder bundling for remaining files
    by_folder: dict[str, list[ScannedFile]] = {}
    for f in files:
        if f.path in used:
            continue
        by_folder.setdefault(f.subject_folder or "_root", []).append(f)

    for folder, ffiles in sorted(by_folder.items()):
        embed_map: dict[str, ScannedFile] = {}
        for f in ffiles:
            if f.role in ("onenote_export", "answer"):
                for ref in f.embedded_refs:
                    embed_map[ref] = f

        for tf in [f for f in ffiles if f.role == "task"]:
            if tf.path in used:
                continue
            b = TaskBundle(
                bundle_id=_make_id(folder, tf.filename),
                title=_title(tf),
                subject=folder,
            )
            b.task_files.append(tf.path)
            used.add(tf.path)

            if tf.filename in embed_map:
                on = embed_map[tf.filename]
                if on.path not in used:
                    (b.answer_files if on.role == "answer" else b.duplicates).append(on.path)
                    used.add(on.path)

            for f in ffiles:
                if (
                    f.role == "duplicate"
                    and f.path not in used
                    and f.duplicate_of
                    and Path(f.duplicate_of).name == tf.filename
                ):
                    b.duplicates.append(f.path)
                    used.add(f.path)

            ts = _normalize_stem(tf.filename)
            for f in ffiles:
                if f.path not in used and f.role == "info" and _stems_related(ts, _normalize_stem(f.filename)):
                    b.info_files.append(f.path)
                    used.add(f.path)
            bundles.append(b)

        for f in ffiles:
            if f.path in used or f.role in ("empty", "duplicate"):
                continue
            b = TaskBundle(
                bundle_id=_make_id(folder, f.filename),
                title=_title(f),
                subject=folder,
            )
            if f.role == "task":
                b.task_files.append(f.path)
            elif f.role == "answer":
                b.answer_files.append(f.path)
            elif f.role == "onenote_export":
                b.duplicates.append(f.path)
            else:
                b.info_files.append(f.path)
            used.add(f.path)
            bundles.append(b)

    return bundles


# ============================================================
# Output
# ============================================================


def to_manifest(result: ScanResult) -> dict:
    """Convert scan result to a JSON-serialisable manifest dict."""
    rc: dict[str, int] = {}
    for f in result.files:
        rc[f.role] = rc.get(f.role, 0) + 1

    m: dict = {
        "scan_root": result.scan_root,
        "total_files": result.total_files,
        "subject_folders": result.subject_folders,
        "summary": rc,
        "bundles": [],
        "files": [],
    }
    if result.warnings:
        m["warnings"] = result.warnings

    for b in result.bundles:
        e: dict = {"bundle_id": b.bundle_id, "title": b.title, "subject": b.subject}
        if b.task_files:
            e["task_files"] = b.task_files
        if b.info_files:
            e["info_files"] = b.info_files
        if b.answer_files:
            e["answer_files"] = b.answer_files
        if b.duplicates:
            e["duplicates"] = b.duplicates
        m["bundles"].append(e)

    for f in result.files:
        e = {
            "path": f.path,
            "role": f.role,
            "confidence": round(f.confidence, 2),
            "content_type": f.content_type,
            "size_bytes": f.size_bytes,
        }
        if f.bundle_id:
            e["bundle_id"] = f.bundle_id
        if f.task_signals:
            e["task_signals"] = [p.replace("\\b", "").replace("\\", "")[:25] for p in f.task_signals[:5]]
        if f.info_signals:
            e["info_signals"] = [p.replace("\\b", "").replace("\\", "")[:25] for p in f.info_signals[:5]]
        if f.embedded_refs:
            e["embedded_refs"] = f.embedded_refs
        if f.duplicate_of:
            e["duplicate_of"] = f.duplicate_of
        if f.has_tables:
            e["tables"] = f.table_count
            e["empty_cells"] = f.empty_table_cells
        if f.has_images:
            e["has_images"] = True
        if f.notes:
            e["notes"] = f.notes
        m["files"].append(e)
    return m


def to_summary(result: ScanResult) -> str:
    """Generate a human-readable summary of the scan result."""
    lines = [
        f"Scan: {result.scan_root}",
        f"  {result.total_files} files, {len(result.subject_folders)} folders",
        "",
    ]

    # Role counts
    rc: dict[str, int] = {}
    for f in result.files:
        rc[f.role] = rc.get(f.role, 0) + 1
    icons = {
        "task": "TASK",
        "info": "INFO",
        "answer": "ANSW",
        "onenote_export": "NOTE",
        "duplicate": "DUPL",
        "empty": "EMPT",
        "resource": "RSRC",
        "unknown": "????",
    }
    lines.append("Classification:")
    for r, c in sorted(rc.items(), key=lambda x: -x[1]):
        lines.append(f"  [{icons.get(r, '--'):>4}] {r:.<20} {c}")

    # File type statistics
    type_counts: dict[str, int] = {}
    for f in result.files:
        type_counts[f.content_type] = type_counts.get(f.content_type, 0) + 1
    type_parts = [f"{c} {t}" for t, c in sorted(type_counts.items(), key=lambda x: -x[1])]
    lines.append(f"\nFile types: {', '.join(type_parts)}")

    # Low confidence
    low = [f for f in result.files if f.confidence < 0.5 and f.role != "empty"]
    if low:
        lines.append(f"\n  {len(low)} low-confidence (<50%):")
        for f in low:
            lines.append(f"    {f.path} -> {f.role} ({f.confidence:.0%})")

    # Bundles
    lines.append(f"\nBundles ({len(result.bundles)}):")
    for b in result.bundles:
        p = []
        if b.task_files:
            p.append(f"{len(b.task_files)}T")
        if b.info_files:
            p.append(f"{len(b.info_files)}I")
        if b.answer_files:
            p.append(f"{len(b.answer_files)}A")
        if b.duplicates:
            p.append(f"{len(b.duplicates)}D")
        lines.append(f"  [{b.subject:.<15}] ({'+'.join(p):>8}) {b.title[:55]}")

    # Per-file
    lines.append("\nPer-file:")
    for f in result.files:
        tag = icons.get(f.role, "--")
        extra = ""
        if f.task_signals:
            extra += f" sig={[p.replace(chr(92) + 'b', '')[:15] for p in f.task_signals[:3]]}"
        if f.embedded_refs:
            extra += f" emb={f.embedded_refs}"
        if f.duplicate_of:
            extra += f" dup={Path(f.duplicate_of).name}"
        if f.notes:
            extra += f" [{f.notes}]"
        lines.append(f"  [{tag:>4}] {f.confidence:>4.0%} {f.path}{extra}")

    return "\n".join(lines)


# ============================================================
# Scan Runner
# ============================================================


def scan_directory(root: str | Path) -> ScanResult:
    """Scan a directory tree and return classified files + bundles."""
    root = Path(root)
    result = ScanResult(scan_root=str(root))
    result.subject_folders = [
        d.name for d in sorted(root.iterdir()) if d.is_dir() and not d.name.startswith(("_", "."))
    ]

    all_files: list[Path] = []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if not d.startswith(("_", ".", "__"))]
        for fname in sorted(filenames):
            if not fname.startswith(("_", ".", "~")):
                all_files.append(Path(dirpath) / fname)

    result.total_files = len(all_files)
    for fpath in all_files:
        try:
            sf = classify_file(fpath, root)
            result.files.append(sf)
        except Exception as e:
            logger.warning("Failed to classify %s: %s", fpath, e)
            result.warnings.append(f"Failed: {fpath}: {e}")

    result.bundles = build_bundles(result.files, scan_root=root)
    for b in result.bundles:
        for p in b.task_files + b.info_files + b.answer_files + b.duplicates:
            for f in result.files:
                if f.path == p:
                    f.bundle_id = b.bundle_id

    return result


# ============================================================
# CLI entry point
# ============================================================


def main(argv: list[str] | None = None) -> None:
    """Run the scanner from the command line."""
    import argparse

    parser = argparse.ArgumentParser(description="Scan school document directories")
    parser.add_argument("path", help="Directory to scan")
    parser.add_argument("--verbose", "-v", action="store_true", help="Show per-file structure details")
    parser.add_argument("--json", dest="json_out", metavar="FILE", help="Write JSON manifest to FILE")
    args = parser.parse_args(argv)

    result = scan_directory(args.path)
    print(to_summary(result))

    if args.verbose:
        print("\n--- Verbose structure details ---")
        for f in result.files:
            print(f"\n{f.path} [{f.content_type}]:")
            # Show structure from ContentResult if we re-read
            reader = READERS.get(f.content_type)
            if reader:
                cr = reader(Path(result.scan_root) / f.path)
                if cr.structure:
                    for k, v in cr.structure.items():
                        print(f"  {k}: {v}")

    if args.json_out:
        manifest = to_manifest(result)
        with open(args.json_out, "w", encoding="utf-8") as fh:
            json.dump(manifest, fh, indent=2, ensure_ascii=False)
        print(f"\nJSON manifest written to {args.json_out}")


if __name__ == "__main__":
    main()
