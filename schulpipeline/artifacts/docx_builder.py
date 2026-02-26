"""DOCX artifact builder using python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..styles import (
    VisualStyle, VisualSlotConfig,
    DEFAULT_STYLE, DISABLED_VISUAL_SLOTS,
    _to_docx_rgb,
)

# Visual type icons — text labels for reliable rendering across all fonts
_TYPE_ICONS = {
    "diagram": "[Diagramm]",
    "photo": "[Foto]",
    "chart": "[Grafik]",
    "icon": "[Icon]",
    "screenshot": "[Screenshot]",
}

# Headings that indicate a sources section
_SOURCES_KEYWORDS = frozenset({
    "quellen", "quellenangaben", "referenzen",
    "literaturverzeichnis", "literatur", "sources", "references",
})


def build_docx(
    synthesis: dict[str, Any],
    output_path: Path,
    visual_style: VisualStyle | None = None,
    visual_config: VisualSlotConfig | None = None,
) -> None:
    """Build a DOCX document from synthesis data."""
    if visual_style is None:
        visual_style = DEFAULT_STYLE.visual
    if visual_config is None:
        visual_config = DISABLED_VISUAL_SLOTS

    doc = Document()

    # Set default font from style
    style = doc.styles["Normal"]
    font = style.font
    font.name = visual_style.fonts.body_family
    font.size = Pt(visual_style.fonts.body_size_pt)

    # Apply line spacing
    style.paragraph_format.line_spacing = visual_style.layout.line_spacing

    # Apply margins
    for sec in doc.sections:
        sec.left_margin = Inches(visual_style.layout.margin_inches)
        sec.right_margin = Inches(visual_style.layout.margin_inches)

    # Title
    title = synthesis.get("title", "Dokument")
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = _to_docx_rgb(visual_style.colors.primary)

    # Sections — skip sources sections (handled separately at the end)
    all_sources: list[str] = list(synthesis.get("sources", []))

    for section in synthesis.get("sections", []):
        heading_text = section.get("heading", "")
        content = section.get("content", "")

        # Collect sources from inline sections, don't render them twice
        if heading_text and heading_text.strip().lower() in _SOURCES_KEYWORDS:
            all_sources.extend(section.get("bullet_points", []))
            continue

        if heading_text:
            h = doc.add_heading(heading_text, level=1)
            for run in h.runs:
                run.font.color.rgb = _to_docx_rgb(visual_style.colors.primary)

        if content:
            doc.add_paragraph(content)

        # Add bullet points if present
        for bullet in section.get("bullet_points", []):
            doc.add_paragraph(bullet, style="List Bullet")

        # Visual placeholders
        visuals = section.get("visuals", [])
        if visual_config.enabled and visuals:
            for visual in visuals[:visual_config.max_per_slide]:
                _add_visual_placeholder(doc, visual, visual_style)

    # Deduplicated sources block
    if all_sources:
        # Deduplicate while preserving order
        seen: set[str] = set()
        unique_sources: list[str] = []
        for s in all_sources:
            if s not in seen:
                seen.add(s)
                unique_sources.append(s)

        h = doc.add_heading("Quellen", level=1)
        for run in h.runs:
            run.font.color.rgb = _to_docx_rgb(visual_style.colors.primary)
        for source in unique_sources:
            doc.add_paragraph(source, style="List Bullet")

    doc.save(str(output_path))


def _add_visual_placeholder(
    doc: Document, visual: dict, style: VisualStyle,
) -> None:
    """Add a gray italic placeholder paragraph for a visual element."""
    icon = _TYPE_ICONS.get(visual.get("type", ""), "[Bild]")
    intent = visual.get("intent", "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(f"{icon}  [{intent}]")
    run.font.size = Pt(10)
    run.font.color.rgb = _to_docx_rgb(style.colors.text_muted)
    run.font.italic = True
