"""Worksheet processor — handles messy, unstructured school assignments.

Real school assignments look nothing like clean prompts. They contain:
  - Multiple tasks (Übung 1, 2, 3...) with sub-tasks (a, b, c...)
  - Tables to fill in (Gewinnverteilung, Kostenrechnung, etc.)
  - External links (quiz sites, law references)
  - Legal/source text embedded in the task
  - Missing context (images that weren't pasted, references to handouts)
  - Mixed task types in one document (calculations + essays + links)

This module sits between Intake and Plan. It:
  1. Decomposes a raw document into individual tasks
  2. Classifies each task (calculation, table_fill, text_answer, external_link, unknown)
  3. Solves each task independently
  4. Produces a unified output document matching the original structure

The key insight: the OUTPUT should mirror the INPUT structure.
If the teacher gave a table, the answer is a filled table.
If the teacher asked a question, the answer is under the question.
"""

from __future__ import annotations

import json
from dataclasses import dataclass, field
from typing import Any

from .stages.base import BaseStage
from .stages.intake import _parse_json_response

# ============================================================
# Task Model
# ============================================================

@dataclass
class ParsedTask:
    """A single task extracted from a worksheet."""
    id: str                              # "uebung_2_aufgabe_1a"
    label: str                           # "Übung 2, Aufgabe 1a"
    task_type: str                       # calculation | table_fill | text_answer | external_link | multi_choice | unknown
    text: str                            # The task text
    context: str = ""                    # Supporting text (law references, formulas, etc.)
    table_structure: dict | None = None  # For table_fill: {"headers": [...], "rows": [...]}
    data: dict[str, Any] = field(default_factory=dict)  # Extracted numbers, names, etc.
    external_url: str = ""               # If task references an external site
    solvable: bool = True                # False if we can't solve it (missing image, external quiz)
    skip_reason: str = ""                # Why it's not solvable


@dataclass
class SolvedTask:
    """A task with its solution."""
    task: ParsedTask
    answer: str = ""                     # Text answer
    table_data: dict | None = None       # Filled table: {"headers": [...], "rows": [[...]]}
    calculation_steps: list[str] = field(default_factory=list)  # Show your work
    confidence: float = 1.0              # 0-1, how confident in the answer


@dataclass
class WorksheetResult:
    """Complete worksheet with all tasks solved."""
    title: str
    subject: str
    tasks: list[SolvedTask]
    unsolvable: list[ParsedTask]         # Tasks we couldn't solve
    metadata: dict[str, Any] = field(default_factory=dict)


# ============================================================
# Decompose Stage — parse messy input into structured tasks
# ============================================================

DECOMPOSE_PROMPT = """\
Du bist ein Aufgaben-Parser für Schularbeitsblätter. Analysiere das folgende Dokument
und zerlege es in einzelne Aufgaben.

Antworte mit validem JSON:
{
  "title": "Erkannter Titel oder Thema des Arbeitsblattes",
  "subject": "Erkanntes Fach",
  "tasks": [
    {
      "id": "aufgabe_1",
      "label": "Übung 2, Aufgabe 1a",
      "task_type": "calculation | table_fill | text_answer | external_link | multi_choice | unknown",
      "text": "Der vollständige Aufgabentext",
      "context": "Zusätzlicher Kontext (Gesetzestexte, Formeln, Hinweise)",
      "table_structure": {
        "headers": ["Name", "Kapitalanteil", "4% Zinsen", "Restgewinn", "Gesamtgewinn"],
        "rows": [
          {"label": "Bauer", "values": ["100.000,00 €", "", "", ""]},
          {"label": "Müller", "values": ["150.000,00 €", "", "", ""]},
          {"label": "Summe", "values": ["", "", "", ""]}
        ]
      },
      "data": {
        "gewinn": 130000,
        "gesellschafter": [
          {"name": "Bauer", "einlage": 100000},
          {"name": "Müller", "einlage": 150000}
        ],
        "zinssatz": 0.04
      },
      "external_url": "",
      "solvable": true,
      "skip_reason": ""
    }
  ]
}

Regeln:
- Jede einzelne Teilaufgabe wird eine eigene Task (1a, 1b, 2, 3 sind separate Tasks)
- task_type erkennen:
  - "calculation": Rechenaufgabe mit konkreten Zahlen
  - "table_fill": Tabelle die ausgefüllt werden muss
  - "text_answer": Freitextantwort gefragt
  - "external_link": Verweis auf externe Website/Quiz (nicht lösbar)
  - "multi_choice": Multiple-Choice-Frage
  - "unknown": Nicht erkennbar (z.B. fehlende Bilder)
- table_structure: NUR wenn eine Tabelle erkennbar ist. Headers + Rows mit leeren Feldern.
- data: Alle relevanten Zahlen/Daten aus dem Aufgabentext extrahieren
- context: Gesetzestexte, Formeln, Definitionen die zur Lösung nötig sind
- solvable: false wenn wir die Aufgabe nicht lösen können (externe Links, fehlende Bilder)
- skip_reason: Warum die Aufgabe nicht lösbar ist
- Antworte NUR mit JSON
"""


class DecomposeStage(BaseStage):
    """Decomposes a messy document into individual structured tasks."""

    name = "decompose"
    spec_path = "specs/decompose.json"

    async def execute(self, context: dict[str, Any], backend: Any, config: Any) -> dict[str, Any]:
        intake = context["intake"]
        preset = context.get("preset")

        prompt = DECOMPOSE_PROMPT
        if preset:
            prompt += f"\n\nFachkontext: {preset.system_context}"

        messages = [
            {"role": "system", "content": prompt},
            {"role": "user", "content": f"Arbeitsblatt:\n\n{intake['task_text']}"},
        ]

        response = await backend.complete(
            stage=self.name,
            messages=messages,
            temperature=0.1,
            max_tokens=8192,
            response_format={"type": "json_object"},
        )

        data = _parse_json_response(response.content)
        return data


# ============================================================
# Solve Stage — solves each decomposed task individually
# ============================================================

SOLVE_PROMPT = """\
Du bist ein Aufgaben-Löser für Schulaufgaben. Löse die folgende Aufgabe.

{task_context}

Antworte mit validem JSON:
{{
  "answer": "Textuelle Antwort oder Zusammenfassung",
  "table_data": {{
    "headers": ["Name", "Kapitalanteil", "4% Zinsen", "Restgewinn", "Gesamtgewinn"],
    "rows": [
      {{"label": "Bauer", "values": ["100.000,00 €", "4.000,00 €", "21.000,00 €", "25.000,00 €"]}},
      ...
    ]
  }},
  "calculation_steps": [
    "4% von 100.000 = 4.000 €",
    "4% von 150.000 = 6.000 €",
    "Restgewinn: 130.000 - 10.000 = 120.000 €",
    "Pro Kopf: 120.000 / 2 = 60.000 €"
  ],
  "confidence": 0.95
}}

Regeln:
- Bei Rechenaufgaben: calculation_steps zeigen, Ergebnis in answer UND table_data
- Bei Tabellen: table_data mit allen Feldern ausgefüllt, Zahlen im deutschen Format (Punkt als Tausendertrenner, Komma als Dezimaltrenner)
- Bei Textfragen: Direkte, knappe Antwort
- table_data: null wenn keine Tabelle gefragt ist
- confidence: 0-1, wie sicher du dir bist
- Antworte NUR mit JSON
"""


class SolveStage(BaseStage):
    """Solves each task from the decomposed worksheet."""

    name = "solve"
    spec_path = "specs/solve.json"

    async def execute(self, context: dict[str, Any], backend: Any, config: Any) -> dict[str, Any]:
        decomposed = context["decompose"]
        preset = context.get("preset")

        tasks = decomposed.get("tasks", [])
        solved = []
        unsolvable = []

        for task in tasks:
            if not task.get("solvable", True):
                unsolvable.append(task)
                continue

            solution = await self._solve_task(task, backend, preset)
            solved.append({
                "task": task,
                "solution": solution,
            })

        return {
            "title": decomposed.get("title", "Arbeitsblatt"),
            "subject": decomposed.get("subject", ""),
            "solved_tasks": solved,
            "unsolvable_tasks": unsolvable,
        }

    async def _solve_task(self, task: dict, backend: Any, preset: Any) -> dict:
        """Solve a single task."""
        task_context = f"Aufgabe: {task.get('label', '')}\n"
        task_context += f"Typ: {task.get('task_type', 'unknown')}\n"
        task_context += f"Text: {task.get('text', '')}\n"

        if task.get("context"):
            task_context += f"\nKontext/Gesetzestext:\n{task['context']}\n"

        if task.get("table_structure"):
            table = task["table_structure"]
            task_context += "\nTabelle zum Ausfüllen:\n"
            task_context += f"Spalten: {', '.join(table.get('headers', []))}\n"
            for row in table.get("rows", []):
                task_context += f"  {row.get('label', '')}: {row.get('values', [])}\n"

        if task.get("data"):
            task_context += f"\nExtrahierte Daten: {json.dumps(task['data'], ensure_ascii=False)}\n"

        preset_context = ""
        if preset:
            preset_context = f"\nFachkontext: {preset.system_context}"

        prompt = SOLVE_PROMPT.replace("{task_context}", task_context + preset_context)

        messages = [
            {"role": "system", "content": prompt},
            {"role": "user", "content": f"Löse diese Aufgabe: {task.get('text', '')}"},
        ]

        response = await backend.complete(
            stage="solve",
            messages=messages,
            temperature=0.1,
            max_tokens=4096,
            response_format={"type": "json_object"},
        )

        try:
            return _parse_json_response(response.content)
        except ValueError:
            return {
                "answer": response.content,
                "table_data": None,
                "calculation_steps": [],
                "confidence": 0.5,
            }


# ============================================================
# Format Stage — produces the output document
# ============================================================

def format_worksheet_as_md(result: dict[str, Any]) -> str:
    """Format solved worksheet as Markdown — the simple, always-works path."""
    lines = [f"# {result.get('title', 'Arbeitsblatt')}"]
    if result.get("subject"):
        lines.append(f"*Fach: {result['subject']}*\n")

    for entry in result.get("solved_tasks", []):
        task = entry["task"]
        solution = entry["solution"]

        lines.append(f"\n## {task.get('label', 'Aufgabe')}")
        lines.append(f"*{task.get('text', '')[:200]}*\n")

        # Show calculation steps if present
        steps = solution.get("calculation_steps", [])
        if steps:
            lines.append("**Rechenweg:**")
            for step in steps:
                lines.append(f"- {step}")
            lines.append("")

        # Show table if present
        table = solution.get("table_data")
        if table and table.get("headers"):
            headers = table["headers"]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in table.get("rows", []):
                label = row.get("label", "")
                values = row.get("values", [])
                all_cells = [label] + values if label else values
                # Pad if needed
                while len(all_cells) < len(headers):
                    all_cells.append("")
                lines.append("| " + " | ".join(str(c) for c in all_cells[:len(headers)]) + " |")
            lines.append("")

        # Show text answer
        answer = solution.get("answer", "")
        if answer:
            lines.append(f"**Antwort:** {answer}\n")

    # Unsolvable tasks
    unsolvable = result.get("unsolvable_tasks", [])
    if unsolvable:
        lines.append("\n---\n## Nicht bearbeitbare Aufgaben\n")
        for task in unsolvable:
            reason = task.get("skip_reason", "Unbekannt")
            lines.append(f"- **{task.get('label', '?')}**: {reason}")
            if task.get("external_url"):
                lines.append(f"  Link: {task['external_url']}")

    return "\n".join(lines)


def format_worksheet_as_docx(result: dict[str, Any], output_path) -> None:
    """Format solved worksheet as DOCX with proper tables."""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = result.get("title", "Arbeitsblatt")
    doc.add_heading(title, level=0)

    if result.get("subject"):
        p = doc.add_paragraph()
        run = p.add_run(f"Fach: {result['subject']}")
        run.italic = True

    for entry in result.get("solved_tasks", []):
        task = entry["task"]
        solution = entry["solution"]

        # Task heading
        doc.add_heading(task.get("label", "Aufgabe"), level=2)

        # Task text (italic, as reference)
        p = doc.add_paragraph()
        task_text = task.get("text", "")
        if len(task_text) > 300:
            task_text = task_text[:300] + "..."
        run = p.add_run(task_text)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Calculation steps
        steps = solution.get("calculation_steps", [])
        if steps:
            doc.add_heading("Rechenweg:", level=3)
            for step in steps:
                doc.add_paragraph(step, style="List Bullet")

        # Table
        table_data = solution.get("table_data")
        if table_data and table_data.get("headers"):
            headers = table_data["headers"]
            rows = table_data.get("rows", [])

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"

            # Header row
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = str(header)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)

            # Data rows
            for row_idx, row in enumerate(rows):
                label = row.get("label", "")
                values = row.get("values", [])
                all_cells = [label] + values if label else values

                for col_idx in range(min(len(all_cells), len(headers))):
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell.text = str(all_cells[col_idx])
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

            doc.add_paragraph()  # Spacing after table

        # Text answer
        answer = solution.get("answer", "")
        if answer:
            p = doc.add_paragraph()
            run = p.add_run("Antwort: ")
            run.bold = True
            p.add_run(answer)

    # Unsolvable tasks
    unsolvable = result.get("unsolvable_tasks", [])
    if unsolvable:
        doc.add_heading("Nicht bearbeitbare Aufgaben", level=2)
        for task in unsolvable:
            reason = task.get("skip_reason", "Unbekannt")
            label = task.get("label", "?")
            doc.add_paragraph(f"{label}: {reason}", style="List Bullet")

    doc.save(str(output_path))


# ============================================================
# Output Preset for Worksheets
# ============================================================

WORKSHEET_OUTPUT_CONFIG = {
    "key": "arbeitsblatt",
    "label": "Arbeitsblatt ausfüllen",
    "format": "docx",  # or md
    "stages": ["intake", "decompose", "solve"],  # Custom stage sequence
    "constraints": {
        "show_calculation_steps": True,
        "german_number_format": True,
        "mirror_input_structure": True,
    },
}
