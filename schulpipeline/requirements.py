"""Requirements Report — full documentation before any work begins.

Three-part document that turns vague school assignments into an
auditable, defensible specification:

  Part A: CLASSIFY REPORT
    "This is what we received."
    All requirements extracted from all source documents, structured as
    a numbered list. Each requirement tagged with:
      - Source document
      - Status: clear / ambiguous / contradicted / missing
      - Related constraint (if any)

  Part B: STUDENT AMENDMENTS
    "This is what we decided."
    For every gap, ambiguity, or contradiction found in Part A, the
    student documents their decision. This can be:
      - Auto-generated defaults (pipeline picks the safest option)
      - Manual overrides (student provides their own answer)
    Each amendment references the finding ID it resolves.

  Part C: DEVIATION LOG
    "This is what we can't do, and here's why."
    For every constraint that mathematically or logically can't be met,
    a formal justification:
      - The constraint as stated
      - Why it can't be met (with numbers/evidence)
      - What we'll do instead
      - The impact of the deviation

The key insight: this document exists BEFORE the work starts. Every
decision is traceable. Every deviation is justified. If questioned
later, the student points to the report.

Usage:
  # Generate full report
  schulpipeline run --preset fiae-anforderungen --source Antrag.docx --source Anforderungen.txt

  # Generate report, then manually edit amendments before proceeding
  schulpipeline run --preset fiae-anforderungen --interactive ...
"""

from __future__ import annotations

import json
from dataclasses import dataclass, field
from typing import Any

from .stages.base import BaseStage
from .stages.intake import _parse_json_response


# ============================================================
# Data Model
# ============================================================

@dataclass
class Requirement:
    """A single extracted requirement."""
    id: str                          # "REQ-001"
    text: str                        # The requirement as stated
    source: str                      # Which document
    category: str                    # functional | format | constraint | process | quality
    status: str                      # clear | ambiguous | contradicted | gap
    priority: str                    # must | should | nice_to_have
    related_findings: list[str] = field(default_factory=list)  # Finding IDs from audit
    quote: str = ""                  # Original text from source


@dataclass
class Amendment:
    """A student decision that resolves an audit finding."""
    id: str                          # "AMD-001"
    resolves: str                    # Finding ID (e.g., "F-002") or Requirement ID
    decision: str                    # What was decided
    reasoning: str                   # Why this decision
    source: str                      # "auto" | "manual" | "teacher_confirmed"
    alternatives_considered: list[str] = field(default_factory=list)


@dataclass
class Deviation:
    """A documented deviation from stated constraints."""
    id: str                          # "DEV-001"
    constraint: str                  # The constraint as stated
    constraint_source: str           # Where the constraint comes from
    reason: str                      # Why it can't be met
    evidence: str                    # Numbers, calculations, proof
    alternative: str                 # What we'll do instead
    impact: str                      # What this means for the result
    severity: str                    # minor | moderate | major


# ============================================================
# Part A: Classify Report — extract all requirements
# ============================================================

CLASSIFY_REPORT_PROMPT = """\
Du bist ein Anforderungs-Extraktor für Schulprojekte. Extrahiere ALLE Anforderungen
aus den folgenden Dokumenten als strukturierte Liste.

Antworte mit validem JSON:
{
  "requirements": [
    {
      "id": "REQ-001",
      "text": "Die Anwendung muss ein Login-System haben",
      "source": "Anforderungen.docx",
      "category": "functional | format | constraint | process | quality",
      "status": "clear | ambiguous | contradicted | gap",
      "priority": "must | should | nice_to_have",
      "quote": "Originaltext aus dem Dokument"
    }
  ],
  "implicit_requirements": [
    {
      "id": "IMP-001",
      "text": "Dokumentation in deutscher Sprache",
      "reasoning": "Nicht explizit genannt, aber Standard für deutschsprachigen Unterricht",
      "confidence": 0.9
    }
  ],
  "requirement_count": {
    "total": 15,
    "clear": 10,
    "ambiguous": 3,
    "contradicted": 1,
    "gap": 1
  }
}

Kategorien:
- functional: Was das Ergebnis tun/enthalten muss (Features, Inhalte)
- format: Wie das Ergebnis aussehen muss (Seitenzahl, Layout, Dateiformat)
- constraint: Einschränkungen (Technologie, Umfang, Zeit)
- process: Wie gearbeitet werden soll (Einzelarbeit, Dokumentation)
- quality: Qualitätsanforderungen (Bewertungskriterien, Standards)

Status:
- clear: Eindeutig formuliert, keine Interpretation nötig
- ambiguous: Unklar oder mehrdeutig, Interpretation nötig
- contradicted: Widerspricht einer anderen Anforderung
- gap: Referenziert etwas das nicht bereitgestellt wurde

Regeln:
- JEDE Anforderung einzeln auflisten, auch wenn mehrere im selben Satz stehen
- Implizite Anforderungen separat mit Begründung und Konfidenz
- Wörtliches Zitat aus dem Originaldokument immer angeben
- status "gap" wenn eine Information referenziert wird die fehlt
- Antworte NUR mit JSON
"""


class ClassifyReportStage(BaseStage):
    """Extracts and structures all requirements from source documents."""

    name = "classify_report"
    spec_path = "specs/classify_report.json"

    async def execute(self, context: dict[str, Any], backend: Any, config: Any) -> dict[str, Any]:
        classified = context.get("classify_docs", {})
        audit = context.get("audit", {})
        preset = context.get("preset")
        documents = classified.get("documents", [])

        # Build document content for LLM
        doc_texts = []
        for doc in documents:
            role = doc.get("role", "unknown")
            filename = doc.get("filename", "?")
            content = doc.get("content", "")
            info = doc.get("extracted_info", {})
            doc_texts.append(
                f"--- {filename} (Rolle: {role}) ---\n"
                f"{content}\n"
                f"Extrahierte Infos: {json.dumps(info, ensure_ascii=False)}\n"
            )

        messages = [
            {"role": "system", "content": CLASSIFY_REPORT_PROMPT},
            {"role": "user", "content": f"Extrahiere alle Anforderungen:\n\n{''.join(doc_texts)}"},
        ]

        response = await backend.complete(
            stage=self.name,
            messages=messages,
            temperature=0.1,
            max_tokens=8192,
            response_format={"type": "json_object"},
        )

        data = _parse_json_response(response.content)

        # Cross-reference with audit findings
        requirements = data.get("requirements", [])
        findings = audit.get("findings", [])
        requirements = _cross_reference(requirements, findings)

        data["requirements"] = requirements
        return data


def _cross_reference(requirements: list[dict], findings: list[dict]) -> list[dict]:
    """Link requirements to audit findings that affect them."""
    for req in requirements:
        req_text = req.get("text", "").lower()
        related = []
        for f in findings:
            finding_text = (f.get("title", "") + " " + f.get("detail", "")).lower()
            # Simple keyword overlap check
            req_words = set(req_text.split())
            finding_words = set(finding_text.split())
            overlap = req_words & finding_words - {"die", "der", "das", "und", "oder", "ein", "eine", "mit", "von", "zu"}
            if len(overlap) >= 2:
                related.append(f.get("id", "?"))
        req["related_findings"] = related
    return requirements


# ============================================================
# Part B: Student Amendments — auto-generate decisions
# ============================================================

AMENDMENTS_PROMPT = """\
Du bist ein Entscheidungs-Assistent für Schulprojekte. Für jedes Problem in den
Vorgaben (Lücke, Widerspruch, Mehrdeutigkeit) musst du eine begründete Entscheidung
vorschlagen.

Probleme:
{findings}

Verfügbare Anforderungen:
{requirements}

Antworte mit validem JSON:
{{
  "amendments": [
    {{
      "id": "AMD-001",
      "resolves": "F-002",
      "finding_title": "Widersprüchliche DB-Vorgabe",
      "decision": "SQLite wird verwendet",
      "reasoning": "Die schriftliche Vorgabe (Anforderungen.docx) hat Vorrang vor mündlichen Aussagen. SQLite ist zudem für den Projektumfang ausreichend und einfacher zu deployen.",
      "source": "auto",
      "alternatives_considered": [
        "PostgreSQL (mündlich empfohlen, aber nicht schriftlich dokumentiert)",
        "MySQL (nicht erwähnt, aber ebenfalls möglich)"
      ]
    }}
  ],
  "unresolvable": [
    {{
      "finding_id": "F-005",
      "reason": "Kann nicht automatisch entschieden werden — erfordert Rücksprache mit dem Lehrer"
    }}
  ]
}}

Entscheidungs-Prinzipien (in dieser Reihenfolge):
1. Schriftlich dokumentiert > mündlich ausgesagt
2. Spezifisch > allgemein
3. Neueres Dokument > älteres Dokument
4. Einfachere Lösung > komplexere Lösung (bei gleicher Erfüllung)
5. Industriestandard > Eigenentwicklung

Regeln:
- Für JEDEN Finding mit severity "blocker" oder "warning" eine Entscheidung
- reasoning muss nachvollziehbar und sachlich sein
- alternatives_considered zeigt dass andere Optionen bedacht wurden
- source ist immer "auto" (kann vom Schüler zu "manual" geändert werden)
- Antworte NUR mit JSON
"""


class AmendmentsStage(BaseStage):
    """Generates default decisions for all audit findings."""

    name = "amendments"
    spec_path = "specs/amendments.json"

    async def execute(self, context: dict[str, Any], backend: Any, config: Any) -> dict[str, Any]:
        audit = context.get("audit", {})
        classify_report = context.get("classify_report", {})
        preset = context.get("preset")

        findings = audit.get("findings", [])
        requirements = classify_report.get("requirements", [])

        # Only generate amendments for actionable findings
        actionable = [f for f in findings if f.get("severity") in ("blocker", "warning")]

        if not actionable:
            return {
                "amendments": [],
                "unresolvable": [],
                "all_clear": True,
            }

        findings_text = json.dumps(actionable, ensure_ascii=False, indent=2)[:4000]
        req_text = json.dumps(requirements[:20], ensure_ascii=False, indent=2)[:3000]

        prompt = AMENDMENTS_PROMPT.replace("{findings}", findings_text).replace("{requirements}", req_text)

        messages = [
            {"role": "system", "content": prompt},
            {"role": "user", "content": "Erstelle Entscheidungen für alle offenen Punkte."},
        ]

        response = await backend.complete(
            stage=self.name,
            messages=messages,
            temperature=0.2,
            max_tokens=4096,
            response_format={"type": "json_object"},
        )

        data = _parse_json_response(response.content)
        data["all_clear"] = False
        return data


# ============================================================
# Part C: Deviation Log — justify every deviation
# ============================================================

def generate_deviations(
    audit: dict[str, Any],
    amendments: dict[str, Any],
) -> list[dict]:
    """Generate deviation entries for constraints that can't be met.

    This is mostly deterministic: if the audit found an impossibility,
    and the amendment resolves it with an alternative, that's a deviation.
    """
    deviations = []
    findings = audit.get("findings", [])
    amendment_map = {a.get("resolves"): a for a in amendments.get("amendments", [])}

    dev_counter = 0
    for f in findings:
        if f.get("category") != "impossibility":
            continue

        dev_counter += 1
        finding_id = f.get("id", "?")
        amendment = amendment_map.get(finding_id, {})

        deviation = {
            "id": f"DEV-{dev_counter:03d}",
            "constraint": f.get("title", "?"),
            "constraint_source": ", ".join(f.get("sources", ["?"])),
            "reason": f.get("detail", ""),
            "evidence": _extract_evidence(f),
            "alternative": amendment.get("decision", "Noch keine Entscheidung getroffen"),
            "impact": _assess_impact(f, amendment),
            "severity": _deviation_severity(f),
        }
        deviations.append(deviation)

    return deviations


def _extract_evidence(finding: dict) -> str:
    """Extract numerical/factual evidence from a finding."""
    detail = finding.get("detail", "")
    # Look for numbers in the detail text
    import re
    numbers = re.findall(r'\d+[\.\d]*', detail)
    if numbers:
        return detail  # The detail itself contains the evidence
    return "Siehe Detailbeschreibung im Audit-Bericht"


def _assess_impact(finding: dict, amendment: dict) -> str:
    """Assess the impact of deviating from a constraint."""
    if amendment.get("decision"):
        return (
            f"Statt der ursprünglichen Vorgabe wird folgende Alternative umgesetzt: "
            f"{amendment['decision']}. "
            f"Begründung: {amendment.get('reasoning', 'Siehe Amendment')}."
        )
    return "Auswirkung muss nach Klärung mit dem Lehrer bewertet werden."


def _deviation_severity(finding: dict) -> str:
    """Map finding severity to deviation severity."""
    severity = finding.get("severity", "info")
    return {"blocker": "major", "warning": "moderate", "info": "minor"}.get(severity, "minor")


# ============================================================
# Combined Report Generator
# ============================================================

def build_full_report(
    classify_report: dict[str, Any],
    audit: dict[str, Any],
    amendments: dict[str, Any],
) -> dict[str, Any]:
    """Assemble the three-part requirements report."""
    deviations = generate_deviations(audit, amendments)

    return {
        "part_a": {
            "title": "Teil A: Anforderungsanalyse",
            "subtitle": "Extrahierte Anforderungen aus allen bereitgestellten Dokumenten",
            "requirements": classify_report.get("requirements", []),
            "implicit_requirements": classify_report.get("implicit_requirements", []),
            "counts": classify_report.get("requirement_count", {}),
        },
        "part_b": {
            "title": "Teil B: Eigene Festlegungen",
            "subtitle": "Entscheidungen zu offenen Punkten, Lücken und Widersprüchen",
            "amendments": amendments.get("amendments", []),
            "unresolvable": amendments.get("unresolvable", []),
            "decision_principles": [
                "Schriftlich dokumentiert hat Vorrang vor mündlichen Aussagen",
                "Spezifische Angaben haben Vorrang vor allgemeinen",
                "Bei gleichwertigen Optionen wird die einfachere Lösung gewählt",
                "Industriestandards werden bevorzugt",
            ],
        },
        "part_c": {
            "title": "Teil C: Abweichungsprotokoll",
            "subtitle": "Begründete Abweichungen von nicht einhaltbaren Vorgaben",
            "deviations": deviations,
            "note": (
                "Jede Abweichung ist durch die Analyse in Teil A belegt. "
                "Bei Fragen oder Änderungswünschen kann dieses Dokument als "
                "Gesprächsgrundlage verwendet werden."
            ),
        },
        "audit_summary": audit.get("summary", {}),
    }


# ============================================================
# Markdown Formatter
# ============================================================

def format_report_as_md(report: dict[str, Any]) -> str:
    """Format the three-part report as Markdown."""
    lines = ["# Anforderungsdokumentation\n"]

    summary = report.get("audit_summary", {})
    if summary:
        lines.append(f"**Gesamtbewertung:** {summary.get('verdict', '?')}")
        lines.append(f"**Vollständigkeit:** {summary.get('completeness_score', 0):.0%}")
        lines.append(f"**Machbarkeit:** {summary.get('feasibility_score', 0):.0%}")
        lines.append("")

    # --- Part A ---
    part_a = report.get("part_a", {})
    lines.append(f"## {part_a.get('title', 'Teil A')}")
    lines.append(f"*{part_a.get('subtitle', '')}*\n")

    counts = part_a.get("counts", {})
    if counts:
        lines.append(f"| Status | Anzahl |")
        lines.append(f"|---|---|")
        for status in ["clear", "ambiguous", "contradicted", "gap"]:
            label = {"clear": "Eindeutig", "ambiguous": "Mehrdeutig",
                     "contradicted": "Widersprüchlich", "gap": "Fehlend"}.get(status, status)
            lines.append(f"| {label} | {counts.get(status, 0)} |")
        lines.append("")

    for req in part_a.get("requirements", []):
        status_icon = {
            "clear": "\u2705", "ambiguous": "\u26A0\uFE0F",
            "contradicted": "\u274C", "gap": "\u2753"
        }.get(req.get("status", ""), "")

        lines.append(f"### {req.get('id', '?')}: {req.get('text', '?')} {status_icon}")
        lines.append(f"*Quelle: {req.get('source', '?')} | "
                      f"Kategorie: {req.get('category', '?')} | "
                      f"Priorität: {req.get('priority', '?')}*")
        if req.get("quote"):
            lines.append(f'> "{req["quote"]}"')
        if req.get("related_findings"):
            lines.append(f"Betrifft: {', '.join(req['related_findings'])}")
        lines.append("")

    implicit = part_a.get("implicit_requirements", [])
    if implicit:
        lines.append("### Implizite Anforderungen")
        lines.append("*Nicht explizit genannt, aber aus dem Kontext abgeleitet:*\n")
        for imp in implicit:
            conf = imp.get("confidence", 0)
            lines.append(f"- **{imp.get('id', '?')}**: {imp.get('text', '?')} "
                          f"({conf:.0%} Konfidenz)")
            lines.append(f"  Begründung: {imp.get('reasoning', '?')}")
        lines.append("")

    # --- Part B ---
    part_b = report.get("part_b", {})
    lines.append(f"## {part_b.get('title', 'Teil B')}")
    lines.append(f"*{part_b.get('subtitle', '')}*\n")

    lines.append("**Entscheidungsprinzipien:**")
    for p in part_b.get("decision_principles", []):
        lines.append(f"1. {p}")
    lines.append("")

    for amd in part_b.get("amendments", []):
        lines.append(f"### {amd.get('id', '?')}: {amd.get('finding_title', amd.get('resolves', '?'))}")
        lines.append(f"**Entscheidung:** {amd.get('decision', '?')}")
        lines.append(f"**Begründung:** {amd.get('reasoning', '?')}")
        source_label = {
            "auto": "Automatisch (Pipeline-Empfehlung)",
            "manual": "Manuell (Schüler-Entscheidung)",
            "teacher_confirmed": "Vom Lehrer bestätigt",
        }.get(amd.get("source", "auto"), amd.get("source", "?"))
        lines.append(f"*Quelle: {source_label}*")

        alts = amd.get("alternatives_considered", [])
        if alts:
            lines.append(f"*Verworfene Alternativen:*")
            for alt in alts:
                lines.append(f"  - {alt}")
        lines.append("")

    unresolvable = part_b.get("unresolvable", [])
    if unresolvable:
        lines.append("### Offene Punkte (Rücksprache erforderlich)")
        for u in unresolvable:
            lines.append(f"- **{u.get('finding_id', '?')}**: {u.get('reason', '?')}")
        lines.append("")

    # --- Part C ---
    part_c = report.get("part_c", {})
    lines.append(f"## {part_c.get('title', 'Teil C')}")
    lines.append(f"*{part_c.get('subtitle', '')}*\n")

    deviations = part_c.get("deviations", [])
    if not deviations:
        lines.append("*Keine Abweichungen erforderlich — alle Vorgaben sind einhaltbar.*\n")
    else:
        for dev in deviations:
            severity_icon = {"major": "\U0001F534", "moderate": "\U0001F7E1", "minor": "\U0001F535"}.get(
                dev.get("severity", "minor"), "")
            lines.append(f"### {dev.get('id', '?')}: {dev.get('constraint', '?')} {severity_icon}")
            lines.append(f"**Quelle:** {dev.get('constraint_source', '?')}")
            lines.append(f"**Grund:** {dev.get('reason', '?')}")
            lines.append(f"**Nachweis:** {dev.get('evidence', '?')}")
            lines.append(f"**Alternative:** {dev.get('alternative', '?')}")
            lines.append(f"**Auswirkung:** {dev.get('impact', '?')}")
            lines.append("")

    if part_c.get("note"):
        lines.append(f"---\n*{part_c['note']}*")

    return "\n".join(lines)


# ============================================================
# DOCX Formatter
# ============================================================

def format_report_as_docx(report: dict[str, Any], output_path) -> None:
    """Format the three-part report as DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Anforderungsdokumentation", level=0)

    # Summary box
    summary = report.get("audit_summary", {})
    if summary:
        p = doc.add_paragraph()
        p.add_run("Gesamtbewertung: ").bold = True
        verdict_run = p.add_run(summary.get("verdict", "?"))
        if summary.get("blockers", 0) > 0:
            verdict_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        metrics = [
            ("Vollständigkeit", f"{summary.get('completeness_score', 0):.0%}"),
            ("Machbarkeit", f"{summary.get('feasibility_score', 0):.0%}"),
            ("Feststellungen", str(summary.get("total_findings", 0))),
        ]
        for i, (label, value) in enumerate(metrics):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        doc.add_paragraph()

    # --- Part A ---
    part_a = report.get("part_a", {})
    doc.add_heading(part_a.get("title", "Teil A"), level=1)
    p = doc.add_paragraph()
    run = p.add_run(part_a.get("subtitle", ""))
    run.italic = True

    for req in part_a.get("requirements", []):
        status_label = {
            "clear": "[Eindeutig]", "ambiguous": "[Mehrdeutig]",
            "contradicted": "[Widerspruch]", "gap": "[Fehlend]"
        }.get(req.get("status", ""), "")

        doc.add_heading(
            f"{req.get('id', '?')}: {req.get('text', '?')} {status_label}",
            level=3
        )

        meta = doc.add_paragraph()
        meta.add_run("Quelle: ").bold = True
        meta.add_run(req.get("source", "?"))
        meta.add_run("  |  Kategorie: ").bold = True
        meta.add_run(req.get("category", "?"))
        meta.add_run("  |  Priorität: ").bold = True
        meta.add_run(req.get("priority", "?"))

        if req.get("quote"):
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f'"{req["quote"]}"')
            q_run.italic = True
            q_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Part B ---
    part_b = report.get("part_b", {})
    doc.add_heading(part_b.get("title", "Teil B"), level=1)
    p = doc.add_paragraph()
    run = p.add_run(part_b.get("subtitle", ""))
    run.italic = True

    doc.add_heading("Entscheidungsprinzipien", level=3)
    for principle in part_b.get("decision_principles", []):
        doc.add_paragraph(principle, style="List Number")

    for amd in part_b.get("amendments", []):
        doc.add_heading(
            f"{amd.get('id', '?')}: {amd.get('finding_title', amd.get('resolves', '?'))}",
            level=3
        )
        p = doc.add_paragraph()
        p.add_run("Entscheidung: ").bold = True
        p.add_run(amd.get("decision", "?"))

        p = doc.add_paragraph()
        p.add_run("Begründung: ").bold = True
        p.add_run(amd.get("reasoning", "?"))

        source_label = {
            "auto": "Automatisch (Pipeline-Empfehlung)",
            "manual": "Manuell (Schüler-Entscheidung)",
            "teacher_confirmed": "Vom Lehrer bestätigt",
        }.get(amd.get("source", "auto"), amd.get("source", "?"))
        p = doc.add_paragraph()
        p.add_run("Herkunft: ").bold = True
        p.add_run(source_label)

        alts = amd.get("alternatives_considered", [])
        if alts:
            p = doc.add_paragraph()
            p.add_run("Verworfene Alternativen:").italic = True
            for alt in alts:
                doc.add_paragraph(alt, style="List Bullet")

    # --- Part C ---
    part_c = report.get("part_c", {})
    doc.add_heading(part_c.get("title", "Teil C"), level=1)
    p = doc.add_paragraph()
    run = p.add_run(part_c.get("subtitle", ""))
    run.italic = True

    deviations = part_c.get("deviations", [])
    if not deviations:
        doc.add_paragraph(
            "Keine Abweichungen erforderlich — alle Vorgaben sind einhaltbar."
        ).italic = True
    else:
        for dev in deviations:
            doc.add_heading(f"{dev.get('id', '?')}: {dev.get('constraint', '?')}", level=3)

            fields = [
                ("Quelle", dev.get("constraint_source", "?")),
                ("Grund", dev.get("reason", "?")),
                ("Nachweis", dev.get("evidence", "?")),
                ("Alternative", dev.get("alternative", "?")),
                ("Auswirkung", dev.get("impact", "?")),
                ("Schwere", dev.get("severity", "?")),
            ]
            for label, value in fields:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(value))

    # Footer note
    if part_c.get("note"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(part_c["note"])
        run.italic = True
        run.font.size = Pt(9)

    doc.save(str(output_path))
