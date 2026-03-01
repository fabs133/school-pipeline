"""Vorgaben-Audit — specification validation for school assignments.

Before doing any work, this module answers one question:
  "Are the requirements complete, consistent, and fulfillable?"

This is the most valuable output the pipeline produces. Not the filled-in
forms, not the presentations — the documented proof that the specifications
are contradictory, incomplete, or impossible BEFORE the student starts.

The audit produces a structured report with:
  1. CONTRADICTIONS — Document A says X, Document B says Y
  2. GAPS — Required information that no document provides
  3. IMPOSSIBILITIES — Constraints that conflict (e.g., "one page" + 15 required fields)
  4. AMBIGUITIES — Vague requirements that could be interpreted multiple ways
  5. UNDERDEFINED — References to things that don't exist yet ("see grading rubric" with no rubric)

Each finding includes:
  - Severity (blocker / warning / info)
  - Source documents (which files/statements conflict)
  - Direct quotes from the source material
  - A clear, factual description of the problem

The audit is DETERMINISTIC where possible. Character limits vs. field counts
is pure math. Contradictions between documents are string comparison.
Only semantic ambiguity detection requires LLM assistance.
"""

from __future__ import annotations

import logging
import math
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from .stages.base import BaseStage
from .stages.intake import _parse_json_response

_logger = logging.getLogger("schulpipeline.audit")

# ============================================================
# Audit Model
# ============================================================


@dataclass
class AuditFinding:
    """A single finding from the specification audit."""

    id: str  # "F-001"
    category: str  # contradiction | gap | impossibility | ambiguity | underdefined
    severity: str  # blocker | warning | info
    title: str  # Short description
    detail: str  # Full explanation
    sources: list[str]  # Which documents are involved
    quotes: list[str]  # Direct quotes from source material
    recommendation: str = ""  # What should be clarified/changed


@dataclass
class AuditReport:
    """Complete audit of assignment specifications."""

    title: str
    timestamp: str
    documents_analyzed: list[str]
    findings: list[AuditFinding]
    summary: AuditSummary
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class AuditSummary:
    """Summary statistics for the audit."""

    total_findings: int
    blockers: int
    warnings: int
    info: int
    completeness_score: float  # 0-1, how complete are the specs
    feasibility_score: float  # 0-1, how feasible given constraints
    verdict: str  # "Vorgaben unvollständig" / "Widersprüche gefunden" / "OK"


# ============================================================
# Deterministic Checks — no LLM needed
# ============================================================


def check_template_field_feasibility(
    template: dict,
    all_requirements: list[str],
) -> list[dict]:
    """Check if template fields can physically hold the required content.

    This is pure math: count required fields, estimate content per field,
    check against character limits.
    """
    findings = []
    fields = template.get("fields", []) or []

    if not fields:
        findings.append(
            {
                "category": "gap",
                "severity": "warning",
                "title": f"Keine Felder in Template '{template.get('filename', '?')}' erkannt",
                "detail": "Das Template enthält keine erkennbaren ausfüllbaren Felder. "
                "Entweder ist das Format ungewöhnlich oder es fehlen Platzhalter.",
                "sources": [template.get("filename", "?")],
                "quotes": [],
            }
        )
        return findings

    # Check: do we have enough fields for all requirements?
    req_count = len(all_requirements)
    field_count = len(fields)

    # Check individual field constraints
    for f in fields:
        max_len = f.get("max_length")
        label = f.get("label", "?")

        if max_len and max_len < 20:
            findings.append(
                {
                    "category": "impossibility",
                    "severity": "warning",
                    "title": f"Feld '{label}' hat nur {max_len} Zeichen",
                    "detail": f"Das Feld '{label}' erlaubt maximal {max_len} Zeichen. "
                    f"Das reicht möglicherweise nicht für eine sinnvolle Antwort.",
                    "sources": [template.get("filename", "?")],
                    "quotes": [],
                }
            )

    # Estimate total available space
    total_chars = sum(f.get("max_length", 500) for f in fields)
    avg_chars_per_req = total_chars / max(req_count, 1)

    if req_count > 0 and avg_chars_per_req < 50:
        findings.append(
            {
                "category": "impossibility",
                "severity": "blocker",
                "title": "Template hat nicht genug Platz für alle Anforderungen",
                "detail": f"{req_count} Anforderungen müssen in {field_count} Felder "
                f"mit insgesamt ~{total_chars} Zeichen passen. "
                f"Das ergibt durchschnittlich {avg_chars_per_req:.0f} Zeichen pro Anforderung.",
                "sources": [template.get("filename", "?")],
                "quotes": [],
                "recommendation": "Entweder Anforderungen reduzieren oder mehr Platz im Template schaffen.",
            }
        )

    return findings


def check_page_constraint(
    template: dict,
    max_pages: int | None,
    requirements: list[str],
) -> list[dict]:
    """Check if content can fit within page limits.

    Rough estimate: ~3000 characters per DIN A4 page with normal formatting.
    """
    if not max_pages:
        return []

    findings = []
    chars_per_page = 3000  # Conservative estimate
    max_chars = max_pages * chars_per_page

    # Estimate required content
    fields = template.get("fields", []) or []
    min_required_chars = 0
    for f in fields:
        field_type = f.get("field_type", "text")
        if field_type == "paragraph":
            min_required_chars += 200  # Minimum sensible paragraph
        elif field_type == "text":
            min_required_chars += 30  # Minimum sensible text field
        elif field_type == "table_cell":
            min_required_chars += 15

    # Add overhead for labels, headers, spacing
    overhead_chars = len(fields) * 40  # ~40 chars per field for label + spacing
    total_estimated = min_required_chars + overhead_chars

    if total_estimated > max_chars:
        findings.append(
            {
                "category": "impossibility",
                "severity": "blocker",
                "title": f"Inhalt passt nicht auf {max_pages} Seite(n)",
                "detail": f"Geschätzt werden mindestens {total_estimated} Zeichen benötigt "
                f"({len(fields)} Felder + Labels). "
                f"Bei {max_pages} Seite(n) stehen ca. {max_chars} Zeichen zur Verfügung. "
                f"Differenz: {total_estimated - max_chars} Zeichen zu viel.",
                "sources": [template.get("filename", "?")],
                "quotes": [],
                "recommendation": f"Entweder auf {math.ceil(total_estimated / chars_per_page)} Seiten erweitern "
                f"oder {len(fields) - int(max_chars / (min_required_chars / max(len(fields), 1)))} Felder streichen.",
            }
        )

    return findings


def check_contradictions_deterministic(documents: list[dict]) -> list[dict]:
    """Find contradictions that can be detected by simple comparison.

    Looks for the same topic mentioned in different documents with different values.
    """
    findings = []

    # Collect all extracted info keyed by topic
    topic_sources: dict[str, list[tuple[str, str]]] = {}
    for doc in documents:
        info = doc.get("extracted_info", {})
        filename = doc.get("filename", "?")
        for key, value in info.items():
            if isinstance(value, str):
                topic_sources.setdefault(key, []).append((filename, value))
            elif isinstance(value, list):
                for item in value:
                    if isinstance(item, str):
                        topic_sources.setdefault(key, []).append((filename, item))

    # Check for same topic, different values
    for topic, sources in topic_sources.items():
        if len(sources) < 2:
            continue
        values = set(v for _, v in sources)
        if len(values) > 1:
            source_details = [f'{fn}: "{val}"' for fn, val in sources]
            findings.append(
                {
                    "category": "contradiction",
                    "severity": "warning",
                    "title": f"Widersprüchliche Angaben zu '{topic}'",
                    "detail": "Verschiedene Dokumente machen unterschiedliche Angaben:\n"
                    + "\n".join(f"  - {s}" for s in source_details),
                    "sources": list(set(fn for fn, _ in sources)),
                    "quotes": [v for _, v in sources],
                }
            )

    return findings


def check_missing_references(documents: list[dict]) -> list[dict]:
    """Find references to documents/resources that aren't provided.

    Detects patterns like "siehe Bewertungsschema", "laut Anlage 3", etc.
    """
    findings = []

    # Common reference patterns in German school documents
    reference_markers = [
        "siehe ",
        "laut ",
        "gemäß ",
        "nach ",
        "entsprechend ",
        "Anlage ",
        "Anhang ",
        "Bewertungsschema",
        "Bewertungsbogen",
        "Notenspiegel",
        "Rubrik",
        "Checkliste",
    ]

    provided_filenames = {doc.get("filename", "").lower() for doc in documents}
    all_content = " ".join(doc.get("content", "") for doc in documents if doc.get("role") != "template")

    for marker in reference_markers:
        if marker.lower() in all_content.lower():
            # Check if the referenced thing is actually provided
            # Simple heuristic: if the marker word isn't part of any filename, it's missing
            marker_clean = marker.strip().lower()
            found_in_files = any(marker_clean in fn for fn in provided_filenames)

            if not found_in_files and marker[0].isupper():
                # Find the full sentence containing the marker
                idx = all_content.lower().find(marker.lower())
                context_start = max(0, idx - 30)
                context_end = min(len(all_content), idx + len(marker) + 50)
                context = all_content[context_start:context_end].strip()

                findings.append(
                    {
                        "category": "underdefined",
                        "severity": "warning",
                        "title": f"Referenziertes Dokument nicht bereitgestellt: '{marker.strip()}'",
                        "detail": f"Es wird auf '{marker.strip()}' verwiesen, aber dieses Dokument "
                        f"wurde nicht bereitgestellt.",
                        "sources": ["Alle Dokumente"],
                        "quotes": [f"...{context}..."],
                        "recommendation": f"'{marker.strip()}' anfordern oder klären ob es existiert.",
                    }
                )

    return findings


# ============================================================
# LLM-Assisted Checks — for semantic analysis
# ============================================================

AUDIT_PROMPT = """\
Du bist ein Anforderungs-Auditor für Schulprojekte. Analysiere die folgenden Dokumente
und finde Probleme in den Vorgaben.

Suche nach:
1. WIDERSPRÜCHE: Dokument A sagt X, Dokument B sagt Y (oder selbes Dokument, verschiedene Stellen)
2. LÜCKEN: Wichtige Informationen die fehlen (Abgabedatum? Bewertungskriterien? Technische Vorgaben?)
3. UNMÖGLICHKEITEN: Anforderungen die sich gegenseitig ausschließen
4. MEHRDEUTIGKEITEN: Vorgaben die man auf verschiedene Arten interpretieren kann
5. UNDEFINIERTES: Verweise auf Dinge die nicht existieren oder nicht bereitgestellt wurden

Antworte mit validem JSON:
{
  "findings": [
    {
      "category": "contradiction | gap | impossibility | ambiguity | underdefined",
      "severity": "blocker | warning | info",
      "title": "Kurze Beschreibung",
      "detail": "Ausführliche Erklärung mit konkreten Textverweisen",
      "sources": ["Dokument1.docx", "Mündliche Aussage"],
      "quotes": ["Exakte Zitate aus den Dokumenten"],
      "recommendation": "Was geklärt/geändert werden muss"
    }
  ],
  "missing_information": [
    "Abgabedatum nicht angegeben",
    "Bewertungskriterien fehlen",
    "Technische Umgebung nicht spezifiziert"
  ],
  "completeness_assessment": {
    "score": 0.4,
    "reasoning": "Von 10 üblichen Pflichtangaben fehlen 6"
  }
}

Übliche Pflichtangaben für Schulprojekte die oft fehlen:
- Abgabedatum / Deadline
- Bewertungskriterien / Notenschlüssel
- Erlaubte Hilfsmittel
- Umfang (Seitenzahl, Wortanzahl, Dauer)
- Technische Vorgaben (Sprache, Framework, Datenbank)
- Gruppenarbeit oder Einzelarbeit
- Abgabeformat (digital/print, Dateiformat)
- Präsentationsdauer (falls Präsentation gefordert)
- Ansprechpartner bei Fragen

Regeln:
- NUR echte Probleme melden, keine Spekulationen
- Jedes Finding muss durch konkreten Text belegt sein
- severity "blocker" NUR wenn die Arbeit nicht sinnvoll begonnen werden kann
- Antworte NUR mit JSON
"""


class AuditStage(BaseStage):
    """Audits assignment specifications for completeness and consistency."""

    name = "audit"
    spec_path = "specs/audit.json"

    async def execute(self, context: dict[str, Any], backend: Any, config: Any) -> dict[str, Any]:
        """Executes the method using the provided context, backend, and config.

        :param context: A dictionary containing contextual information.
        :type context: dict[str, Any]
        :param backend: The backend to be used for execution.
        :type backend: Any
        :param config: Configuration settings for the execution.
        :type config: Any
        :return: A dictionary containing the results of the execution.
        :rtype: dict[str, Any]
        """
        classified = context.get("classify_docs", {})
        preset = context.get("preset")
        documents = classified.get("documents", [])

        # Phase 1: Deterministic checks (no LLM needed)
        deterministic_findings = []

        templates = [d for d in documents if d.get("role") == "template"]
        sources = [d for d in documents if d.get("role") == "source"]
        all_requirements = []
        for src in sources:
            all_requirements.extend(src.get("extracted_info", {}).get("requirements", []))

        for tmpl in templates:
            deterministic_findings.extend(check_template_field_feasibility(tmpl, all_requirements))
            # Check page constraints from preset
            max_pages = None
            if preset:
                max_pages = preset.output_constraints.get("max_pages")
            deterministic_findings.extend(check_page_constraint(tmpl, max_pages, all_requirements))

        deterministic_findings.extend(check_contradictions_deterministic(documents))

        # Include contradictions already found by classify_docs
        for c in classified.get("contradictions", []):
            deterministic_findings.append(
                {
                    "category": "contradiction",
                    "severity": "warning",
                    "title": f"Widerspruch: {c.get('topic', '?')}",
                    "detail": f"{c.get('source_a', '?')} vs. {c.get('source_b', '?')}",
                    "sources": [c.get("source_a", "?"), c.get("source_b", "?")],
                    "quotes": [],
                    "recommendation": c.get("recommendation", "Klärung erforderlich"),
                }
            )

        deterministic_findings.extend(check_missing_references(documents))

        # Phase 2: LLM-assisted semantic analysis
        llm_findings = await self._semantic_audit(documents, backend, preset)

        # Merge and deduplicate
        all_findings = self._merge_findings(deterministic_findings, llm_findings)

        # Number findings
        for i, f in enumerate(all_findings):
            f["id"] = f"F-{i + 1:03d}"

        # Calculate summary
        blockers = sum(1 for f in all_findings if f.get("severity") == "blocker")
        warnings = sum(1 for f in all_findings if f.get("severity") == "warning")
        info = sum(1 for f in all_findings if f.get("severity") == "info")

        completeness = llm_findings.get("_completeness_score", 0.5)
        feasibility = 1.0
        if blockers > 0:
            feasibility = max(0.0, 1.0 - (blockers * 0.3))
        if warnings > 0:
            feasibility = max(0.0, feasibility - (warnings * 0.1))

        if blockers > 0:
            verdict = "Vorgaben unvollständig — Arbeit kann nicht sinnvoll begonnen werden"
        elif warnings > 2:
            verdict = "Mehrere Probleme in den Vorgaben — Klärung empfohlen vor Arbeitsbeginn"
        elif warnings > 0:
            verdict = "Kleinere Unklarheiten — Arbeit kann mit Annahmen begonnen werden"
        else:
            verdict = "Vorgaben vollständig und konsistent"

        data = {
            "title": f"Vorgaben-Audit: {classified.get('title', 'Projekt')}",
            "documents_analyzed": [d.get("filename", "?") for d in documents],
            "findings": all_findings,
            "summary": {
                "total_findings": len(all_findings),
                "blockers": blockers,
                "warnings": warnings,
                "info": info,
                "completeness_score": completeness,
                "feasibility_score": round(feasibility, 2),
                "verdict": verdict,
            },
            "missing_information": llm_findings.get("_missing_info", []),
        }

        # Write audit report (moved from pipeline.py post-processing)
        output_dir = context.get("output_dir")
        if output_dir:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            safe_title = "Vorgaben-Audit"
            try:
                audit_docx = output_dir / f"{safe_title}.docx"
                format_audit_as_docx(data, audit_docx)
                # For audit-only mode, this IS the primary output
                is_audit_only = (
                    preset and hasattr(preset, "output_constraints") and preset.output_constraints.get("audit_only")
                )
                if is_audit_only:
                    data["file_path"] = str(audit_docx)
                else:
                    _logger.info(f"Supplementary audit report: {audit_docx}")
            except Exception as e:
                _logger.warning(f"Audit DOCX failed, falling back to MD: {e}")
                audit_md = output_dir / f"{safe_title}.md"
                audit_md.write_text(format_audit_as_md(data), encoding="utf-8")
                is_audit_only = (
                    preset and hasattr(preset, "output_constraints") and preset.output_constraints.get("audit_only")
                )
                if is_audit_only:
                    data["file_path"] = str(audit_md)

        return data

    async def _semantic_audit(self, documents: list[dict], backend: Any, preset: Any) -> dict:
        """Use LLM for semantic analysis of requirements."""
        doc_texts = []
        for doc in documents:
            doc_texts.append(
                f"--- {doc.get('filename', '?')} (Rolle: {doc.get('role', '?')}) ---\n"
                f"{doc.get('content', doc.get('extracted_info', ''))}\n"
            )

        context = "\n\n".join(doc_texts)
        if preset:
            context += f"\n\nVorgabe-Kontext: {preset.system_context}"

        messages = [
            {"role": "system", "content": AUDIT_PROMPT},
            {"role": "user", "content": f"Auditiere diese Projektunterlagen:\n\n{context}"},
        ]

        try:
            response = await backend.complete(
                stage="audit",
                messages=messages,
                temperature=0.1,
                max_tokens=4096,
                response_format={"type": "json_object"},
            )
            data = _parse_json_response(response.content)

            # Extract and return with internal keys for merging
            result = {
                "_findings": data.get("findings", []),
                "_missing_info": data.get("missing_information", []),
                "_completeness_score": data.get("completeness_assessment", {}).get("score", 0.5),
            }
            return result
        except Exception:
            return {"_findings": [], "_missing_info": [], "_completeness_score": 0.5}

    def _merge_findings(self, deterministic: list[dict], llm_result: dict) -> list[dict]:
        """Merge deterministic and LLM findings, removing duplicates."""
        all_findings = list(deterministic)
        llm_findings = llm_result.get("_findings", [])

        # Simple dedup: skip LLM findings whose title is very similar to existing ones
        existing_titles = {f.get("title", "").lower() for f in all_findings}

        for lf in llm_findings:
            title_lower = lf.get("title", "").lower()
            # Skip if there's a very similar title already
            is_duplicate = any(
                title_lower in existing or existing in title_lower for existing in existing_titles if len(existing) > 10
            )
            if not is_duplicate:
                all_findings.append(lf)
                existing_titles.add(title_lower)

        # Sort: blockers first, then warnings, then info
        severity_order = {"blocker": 0, "warning": 1, "info": 2}
        all_findings.sort(key=lambda f: severity_order.get(f.get("severity", "info"), 3))

        return all_findings


# ============================================================
# Report Formatters
# ============================================================


def format_audit_as_md(audit: dict[str, Any]) -> str:
    """Format audit report as Markdown — clear, confrontable, linkable."""
    lines = [f"# {audit.get('title', 'Vorgaben-Audit')}"]
    lines.append("")

    summary = audit.get("summary", {})
    lines.append("## Zusammenfassung")
    lines.append("")
    lines.append(f"**Ergebnis: {summary.get('verdict', '?')}**")
    lines.append("")
    lines.append("| Metrik | Wert |")
    lines.append("|---|---|")
    lines.append(f"| Analysierte Dokumente | {len(audit.get('documents_analyzed', []))} |")
    lines.append(f"| Feststellungen gesamt | {summary.get('total_findings', 0)} |")
    lines.append(f"| Blocker | {summary.get('blockers', 0)} |")
    lines.append(f"| Warnungen | {summary.get('warnings', 0)} |")
    lines.append(f"| Hinweise | {summary.get('info', 0)} |")
    lines.append(f"| Vollständigkeit | {summary.get('completeness_score', 0):.0%} |")
    lines.append(f"| Machbarkeit | {summary.get('feasibility_score', 0):.0%} |")
    lines.append("")

    # Documents analyzed
    docs = audit.get("documents_analyzed", [])
    if docs:
        lines.append("## Analysierte Dokumente")
        lines.append("")
        for d in docs:
            lines.append(f"- {d}")
        lines.append("")

    # Findings by severity
    findings = audit.get("findings", [])
    if not findings:
        lines.append("*Keine Feststellungen — Vorgaben erscheinen vollständig und konsistent.*")
        return "\n".join(lines)

    # Blockers
    blockers = [f for f in findings if f.get("severity") == "blocker"]
    if blockers:
        lines.append("## 🔴 Blocker — Arbeit kann nicht sinnvoll begonnen werden")
        lines.append("")
        for f in blockers:
            lines.extend(_format_finding(f))

    # Warnings
    warns = [f for f in findings if f.get("severity") == "warning"]
    if warns:
        lines.append("## 🟡 Warnungen — Klärung empfohlen")
        lines.append("")
        for f in warns:
            lines.extend(_format_finding(f))

    # Info
    infos = [f for f in findings if f.get("severity") == "info"]
    if infos:
        lines.append("## 🔵 Hinweise")
        lines.append("")
        for f in infos:
            lines.extend(_format_finding(f))

    # Missing information
    missing = audit.get("missing_information", [])
    if missing:
        lines.append("## Fehlende Informationen")
        lines.append("")
        for m in missing:
            lines.append(f"- ❓ {m}")
        lines.append("")

    return "\n".join(lines)


def format_audit_as_docx(audit: dict[str, Any], output_path) -> None:
    """Format audit report as DOCX — for formal submission."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    doc.add_heading(audit.get("title", "Vorgaben-Audit"), level=0)

    # Summary
    summary = audit.get("summary", {})
    doc.add_heading("Zusammenfassung", level=1)

    verdict_para = doc.add_paragraph()
    verdict_run = verdict_para.add_run(f"Ergebnis: {summary.get('verdict', '?')}")
    verdict_run.bold = True
    if summary.get("blockers", 0) > 0:
        verdict_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Summary table
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    summary_data = [
        ("Analysierte Dokumente", str(len(audit.get("documents_analyzed", [])))),
        ("Feststellungen gesamt", str(summary.get("total_findings", 0))),
        ("Blocker", str(summary.get("blockers", 0))),
        ("Warnungen", str(summary.get("warnings", 0))),
        ("Hinweise", str(summary.get("info", 0))),
        ("Vollständigkeit", f"{summary.get('completeness_score', 0):.0%}"),
        ("Machbarkeit", f"{summary.get('feasibility_score', 0):.0%}"),
    ]
    for i, (label, value) in enumerate(summary_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    doc.add_paragraph()

    # Findings
    findings = audit.get("findings", [])
    if findings:
        doc.add_heading("Feststellungen", level=1)

        for f in findings:
            severity = f.get("severity", "info")
            marker = {"blocker": "🔴", "warning": "🟡", "info": "🔵"}.get(severity, "")

            doc.add_heading(f"{marker} {f.get('id', '?')}: {f.get('title', '?')}", level=2)

            # Severity and category
            meta = doc.add_paragraph()
            meta.add_run("Kategorie: ").bold = True
            meta.add_run(f.get("category", "?"))
            meta.add_run("  |  Schwere: ").bold = True
            meta.add_run(severity)

            # Detail
            doc.add_paragraph(f.get("detail", ""))

            # Sources
            sources = f.get("sources", [])
            if sources:
                src_para = doc.add_paragraph()
                src_para.add_run("Quellen: ").bold = True
                src_para.add_run(", ".join(sources))

            # Quotes
            quotes = f.get("quotes", [])
            for q in quotes:
                quote_para = doc.add_paragraph()
                quote_run = quote_para.add_run("  \u201e" + q + "\u201c")
                quote_run.italic = True
                quote_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Recommendation
            rec = f.get("recommendation", "")
            if rec:
                rec_para = doc.add_paragraph()
                rec_para.add_run("Empfehlung: ").bold = True
                rec_para.add_run(rec)

            doc.add_paragraph()  # spacing

    # Missing information
    missing = audit.get("missing_information", [])
    if missing:
        doc.add_heading("Fehlende Informationen", level=1)
        for m in missing:
            doc.add_paragraph(f"❓ {m}", style="List Bullet")

    doc.save(str(output_path))


def _format_finding(f: dict) -> list[str]:
    """Format a single finding as Markdown lines."""
    lines = []
    lines.append(f"### {f.get('id', '?')}: {f.get('title', '?')}")
    lines.append(f"*Kategorie: {f.get('category', '?')}*")
    lines.append("")
    lines.append(f.get("detail", ""))
    lines.append("")

    sources = f.get("sources", [])
    if sources:
        lines.append(f"**Quellen:** {', '.join(sources)}")

    quotes = f.get("quotes", [])
    if quotes:
        for q in quotes:
            lines.append("> \u201e" + q + "\u201c")

    rec = f.get("recommendation", "")
    if rec:
        lines.append(f"**Empfehlung:** {rec}")

    lines.append("")
    return lines
