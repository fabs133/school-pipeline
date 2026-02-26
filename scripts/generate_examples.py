"""Generate synthetic example documents for schulpipeline testing.

All content is fully synthetic — no copyrighted material, no real student data.
Uses python-docx to create .docx files with specific structural patterns
that the scanner recognizes for classification.

Fictional context: Berufskolleg Beispielstadt.

Usage:
    python scripts/generate_examples.py
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent.parent
EXAMPLES = ROOT / "examples"
TASKS = EXAMPLES / "tasks"
DE = TASKS / "DE-BSP"
WI = TASKS / "WI-BSP"


# ── Helpers ──────────────────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row()
        for i, cell in enumerate(row_data):
            row.cells[i].text = cell


def _minimal_pdf(path: Path) -> None:
    """Write a minimal valid PDF file (for .docx.pdf duplicate detection)."""
    pdf = (
        b"%PDF-1.0\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R>>endobj\n"
        b"xref\n0 4\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"trailer<</Size 4/Root 1 0 R>>\n"
        b"startxref\n190\n%%EOF\n"
    )
    path.write_bytes(pdf)


# ── Document generators ─────────────────────────────────────────────

def gen_energie_task(path: Path) -> None:
    """Task doc: imperative-Sie keywords, slide requirement.

    Scanner triggers: _A1_ in filename → +3 task_score,
    strong keywords (Formulieren Sie, Erstellen Sie, Recherchieren Sie) → +6.
    Expected role: task.
    """
    doc = Document()
    _add_heading(doc, "Aufgabe: Erneuerbare Energien")
    _add_para(doc, "Berufskolleg Beispielstadt — Fach Deutsch, Klasse DE-BSP")
    _add_para(doc, "Bearbeitungszeit: 90 Minuten")
    doc.add_paragraph()

    _add_heading(doc, "Ausgangssituation", level=2)
    _add_para(
        doc,
        "In Ihrem Ausbildungsbetrieb sollen die Auszubildenden eine "
        "Praesentation zum Thema erneuerbare Energien vorbereiten. "
        "Ihr Ausbilder hat Ihnen dazu folgende Aufgaben gegeben.",
    )
    doc.add_paragraph()

    _add_heading(doc, "Aufgaben", level=2)
    _add_para(
        doc,
        "Aufgabe 1: Recherchieren Sie die wichtigsten Formen erneuerbarer "
        "Energien in Deutschland (Solar, Wind, Biomasse, Wasserkraft, Geothermie).",
    )
    _add_para(
        doc,
        "Aufgabe 2: Formulieren Sie fuer jede Energieform eine kurze "
        "Definition mit Vor- und Nachteilen.",
    )
    _add_para(
        doc,
        "Aufgabe 3: Erstellen Sie eine Praesentation mit mindestens "
        "8 Folien, die folgende Punkte abdeckt:",
    )

    bullets = [
        "Titelfolie mit Thema und Ihrem Namen",
        "Ueberblick ueber erneuerbare Energien",
        "Solarenergie: Funktionsweise und Einsatzgebiete",
        "Windenergie: Onshore und Offshore im Vergleich",
        "Biomasse und Geothermie als Ergaenzung",
        "Aktuelle Zahlen zum Energiemix in Deutschland",
        "Vor- und Nachteile erneuerbarer Energien",
        "Quellenfolie",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph()
    _add_para(
        doc,
        "Aufgabe 4: Beschreiben Sie in eigenen Worten, warum der Ausbau "
        "erneuerbarer Energien fuer Deutschland wichtig ist.",
    )
    _add_para(
        doc,
        "Aufgabe 5: Nennen Sie mindestens drei Quellen, die Sie fuer "
        "Ihre Recherche verwendet haben.",
    )
    doc.add_paragraph()
    _add_para(
        doc,
        "Hinweis: Achten Sie auf eine uebersichtliche Gestaltung und "
        "verwenden Sie Stichpunkte statt ganzer Saetze auf den Folien.",
    )
    doc.save(str(path))


def gen_energie_texte(path: Path) -> None:
    """Info doc: Fachtext with info keywords, no imperatives.

    Scanner triggers: _Texte in filename → TEXT_MATERIAL +5 info, -3 task,
    also INFO_FILENAME_RE → +3 info.
    Plus info keywords: Grundsaetzlich, Definition, Man unterscheidet → +3.
    Expected role: info.
    """
    doc = Document()
    _add_heading(doc, "Informationsteil: Erneuerbare Energien")
    _add_para(doc, "Berufskolleg Beispielstadt — Begleitmaterial")
    doc.add_paragraph()

    _add_heading(doc, "1. Definition", level=2)
    _add_para(
        doc,
        "Definition: Erneuerbare Energien sind Energiequellen, die sich "
        "auf natuerliche Weise erneuern und daher praktisch unerschoepflich "
        "zur Verfuegung stehen.",
    )
    _add_para(
        doc,
        "Grundsaetzlich unterscheidet man zwischen direkter Nutzung "
        "(z. B. Solarwaerme) und indirekter Nutzung (z. B. Windkraft "
        "als Folge von Sonneneinstrahlung).",
    )
    doc.add_paragraph()

    _add_heading(doc, "2. Formen erneuerbarer Energien", level=2)
    _add_para(
        doc,
        "Man unterscheidet fuenf Hauptformen erneuerbarer Energien, "
        "die in Deutschland genutzt werden:",
    )
    forms = [
        (
            "Solarenergie",
            "Die Sonne liefert in einer Stunde mehr Energie auf die "
            "Erdoberflaeche, als die gesamte Menschheit in einem Jahr "
            "verbraucht. Photovoltaikanlagen wandeln Sonnenlicht direkt "
            "in elektrischen Strom um. Solarthermieanlagen nutzen die "
            "Waerme der Sonne zum Heizen.",
        ),
        (
            "Windenergie",
            "Windkraftanlagen wandeln die kinetische Energie des Windes "
            "in Strom um. In Deutschland stehen ueber 30.000 Windraeder. "
            "Onshore-Anlagen stehen an Land, Offshore-Anlagen auf dem Meer. "
            "Offshore-Anlagen liefern gleichmaessigeren Strom.",
        ),
        (
            "Biomasse",
            "Biomasse umfasst organische Stoffe wie Holz, Pflanzenreste "
            "und Biogas. Biogasanlagen vergaeren organische Abfaelle und "
            "erzeugen dabei Methan, das zur Stromerzeugung genutzt wird.",
        ),
        (
            "Wasserkraft",
            "Wasserkraftwerke nutzen die Stroemung oder das Gefaelle von "
            "Fluessen zur Stromerzeugung. In Deutschland sind etwa 7.300 "
            "Wasserkraftanlagen in Betrieb.",
        ),
        (
            "Geothermie",
            "Geothermie nutzt die Waerme aus dem Erdinneren. In Tiefen "
            "von 3.000 bis 5.000 Metern herrschen Temperaturen von ueber "
            "100 Grad Celsius. Diese Waerme kann zum Heizen oder zur "
            "Stromerzeugung genutzt werden.",
        ),
    ]
    for title, text in forms:
        _add_heading(doc, title, level=3)
        _add_para(doc, text)
        doc.add_paragraph()

    _add_heading(doc, "3. Energiemix in Deutschland", level=2)
    _add_para(
        doc,
        "Zu den wichtigsten Kennzahlen zaehlen der Anteil erneuerbarer "
        "Energien am Bruttostromverbrauch sowie die installierte Leistung. "
        "Der Anteil erneuerbarer Energien am Stromverbrauch lag im "
        "vergangenen Jahr bei ueber 50 Prozent.",
    )
    _add_para(
        doc,
        "Beispiele fuer den Ausbau: Die installierte Leistung von "
        "Photovoltaik hat sich in den letzten zehn Jahren verdreifacht. "
        "Windenergie liefert den groessten Anteil am erneuerbaren Strom.",
    )
    doc.add_paragraph()

    _add_heading(doc, "4. Herausforderungen", level=2)
    _add_para(
        doc,
        "Trotz der Vorteile gibt es Herausforderungen: Speicherung "
        "ueberschuessiger Energie, Netzausbau, Fluktuationen bei "
        "Wind- und Solarstrom sowie Akzeptanz in der Bevoelkerung.",
    )
    _add_para(
        doc,
        "Information: Zum Thema Energiespeicherung gibt es verschiedene "
        "Ansaetze wie Batteriespeicher, Pumpspeicherkraftwerke und die "
        "Umwandlung in Wasserstoff (Power-to-Gas).",
    )
    doc.save(str(path))


def gen_energie_onenote(path: Path) -> None:
    """OneNote export with embed marker, few substantive lines → onenote_export.

    Scanner triggers: filename ends with OneNote.docx → ONENOTE_FILENAME_RE,
    contains <<IT_B1_A1_Energie.docx>> embed.
    Substantive lines <= 3 → role=onenote_export.
    """
    doc = Document()
    _add_para(doc, "<<IT_B1_A1_Energie.docx>>")
    _add_para(doc, "Notizen:")
    _add_para(doc, "x")
    doc.save(str(path))


def gen_bewertungsbogen(path: Path) -> None:
    """Planning/grading doc: table with grading criteria.

    Scanner triggers: Bewertung in filename → PLANNING_FILENAME_RE → +2 info.
    Info keywords inside → more info_score.
    Expected role: info.
    """
    doc = Document()
    _add_heading(doc, "Bewertungsbogen: Praesentation Erneuerbare Energien")
    _add_para(doc, "Berufskolleg Beispielstadt — Bewertungskriterien")
    doc.add_paragraph()

    _add_para(
        doc,
        "Information: Dieser Bewertungsbogen dient der Benotung der "
        "Schueler-Praesentationen zum Thema erneuerbare Energien.",
    )
    doc.add_paragraph()

    _add_table(
        doc,
        headers=["Kriterium", "Punkte (max.)", "Punkte (erreicht)", "Bemerkungen"],
        rows=[
            ["Inhaltliche Richtigkeit", "20", "", ""],
            ["Struktur und Gliederung", "15", "", ""],
            ["Gestaltung der Folien", "15", "", ""],
            ["Quellenangaben", "10", "", ""],
            ["Vortragsstil", "20", "", ""],
            ["Zeitmanagement", "10", "", ""],
            ["Gesamteindruck", "10", "", ""],
        ],
    )
    doc.add_paragraph()
    _add_para(doc, "Gesamtpunktzahl: ___ / 100")
    _add_para(doc, "Note: ___")
    doc.save(str(path))


def gen_angebot_nachfrage_task(path: Path) -> None:
    """Task doc: imperative keywords + fill-in table with empty cells.

    Scanner triggers: Aufgaben in filename → TASK_FILENAME_RE → +3 task,
    strong keywords (Ergaenzen Sie, Erklaeren Sie, Nennen Sie) → +6.
    empty_table_cells > 2 → +3 task.
    Expected role: task.
    """
    doc = Document()
    _add_heading(doc, "Aufgaben: Angebot und Nachfrage")
    _add_para(doc, "Berufskolleg Beispielstadt — Fach Wirtschaft, Klasse WI-BSP")
    _add_para(doc, "Bearbeitungszeit: 45 Minuten")
    doc.add_paragraph()

    _add_heading(doc, "Ausgangssituation", level=2)
    _add_para(
        doc,
        "Ein Unternehmen moechte den Preis fuer ein neues Produkt "
        "festlegen. Dazu muss es die Zusammenhaenge zwischen Angebot, "
        "Nachfrage und Preis verstehen.",
    )
    doc.add_paragraph()

    _add_heading(doc, "Aufgabe 1", level=2)
    _add_para(
        doc,
        "Erklaeren Sie den Unterschied zwischen Angebot und Nachfrage "
        "in eigenen Worten.",
    )
    doc.add_paragraph()

    _add_heading(doc, "Aufgabe 2", level=2)
    _add_para(
        doc,
        "Ergaenzen Sie die folgende Tabelle mit den fehlenden Werten:",
    )
    _add_table(
        doc,
        headers=["Preis (EUR)", "Angebotene Menge", "Nachgefragte Menge", "Marktlage"],
        rows=[
            ["5,00", "100", "500", ""],
            ["10,00", "200", "350", ""],
            ["15,00", "300", "300", ""],
            ["20,00", "400", "", ""],
            ["25,00", "", "", ""],
            ["30,00", "", "", ""],
        ],
    )
    doc.add_paragraph()

    _add_heading(doc, "Aufgabe 3", level=2)
    _add_para(
        doc,
        "Nennen Sie drei Faktoren, die das Angebot beeinflussen, und "
        "drei Faktoren, die die Nachfrage beeinflussen.",
    )
    doc.add_paragraph()

    _add_heading(doc, "Aufgabe 4", level=2)
    _add_para(
        doc,
        "Beschreiben Sie, was passiert, wenn der Preis eines Produkts "
        "ueber dem Gleichgewichtspreis liegt.",
    )
    doc.save(str(path))


def gen_angebot_nachfrage_onenote(path: Path) -> None:
    """OneNote export with embed + >3 substantive student answer lines → answer.

    Scanner triggers: filename ends with OneNote.docx → ONENOTE_FILENAME_RE,
    contains <<Aufgaben_Angebot_Nachfrage.docx>> embed.
    Substantive lines > 3 → role=answer.
    """
    doc = Document()
    _add_para(doc, "<<Aufgaben_Angebot_Nachfrage.docx>>")
    _add_para(doc, "Meine Notizen zu Aufgabe 1:")
    _add_para(
        doc,
        "Angebot ist die Menge an Guetern, die Verkaufer bereit sind "
        "zu einem bestimmten Preis zu verkaufen.",
    )
    _add_para(
        doc,
        "Nachfrage ist die Menge, die Kaeufer zu einem bestimmten "
        "Preis kaufen moechten.",
    )
    _add_para(doc, "Wenn Angebot und Nachfrage gleich sind, ist Gleichgewicht.")
    _add_para(doc, "Bei hohem Preis sinkt die Nachfrage.")
    doc.save(str(path))


# ── Main ─────────────────────────────────────────────────────────────

def main() -> None:
    # Clean and recreate
    if EXAMPLES.exists():
        shutil.rmtree(EXAMPLES)

    for d in (DE, WI, EXAMPLES / "output", EXAMPLES / "manifests"):
        d.mkdir(parents=True, exist_ok=True)

    # DE-BSP documents
    gen_energie_task(DE / "IT_B1_A1_Energie.docx")
    gen_energie_texte(DE / "IT_B1_A1_Energie_Texte.docx")
    _minimal_pdf(DE / "IT_B1_A1_Energie.docx.pdf")
    gen_energie_onenote(DE / "EnergieOneNote.docx")
    gen_bewertungsbogen(DE / "Bewertungsbogen.docx")
    (DE / "Leer.docx").write_bytes(b"")  # empty file

    # WI-BSP documents
    gen_angebot_nachfrage_task(WI / "Aufgaben_Angebot_Nachfrage.docx")
    _minimal_pdf(WI / "Aufgaben_Angebot_Nachfrage.docx.pdf")
    gen_angebot_nachfrage_onenote(WI / "AngebotNachfrageOneNote.docx")

    print(f"Generated {sum(1 for _ in EXAMPLES.rglob('*') if _.is_file())} files in {EXAMPLES}/")
    for p in sorted(EXAMPLES.rglob("*")):
        if p.is_file():
            size = p.stat().st_size
            print(f"  {p.relative_to(EXAMPLES)}  ({size:,} bytes)")


if __name__ == "__main__":
    main()
